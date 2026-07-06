import unittest
from pathlib import Path
import sys
import tempfile

BACKEND_DIR = Path(__file__).resolve().parents[1]
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

from app.services.text_splitter import (
    detect_plain_text_headings,
    PdfPageText,
    flatten_sections_to_blocks,
    is_only_heading_context_section,
    split_markdown_sections,
    split_document_text,
    split_pdf_sections,
    split_plain_text_sections,
)
from app.api.document import extract_text_from_file
from app.services.document_splitter.splitter import split_document_text as split_document_text_pipeline
from app.services.document_splitter.splitter import build_document_sections, parse_splitter_source
from app.services.document_splitter.parsers.pdf_layout_parser import (
    PdfLayoutLine,
    build_pdf_paragraph_text,
    is_probable_pdf_noise_paragraph,
    pdf_layout_document_to_text,
)


class TextSplitterStructureTests(unittest.TestCase):
    def create_temp_workbook(self, build_callback):
        try:
            from openpyxl import Workbook
        except ModuleNotFoundError:
            self.skipTest("openpyxl is not installed")

        with tempfile.TemporaryDirectory() as temp_dir:
            workbook = Workbook()
            build_callback(workbook)
            workbook_path = Path(temp_dir) / "test.xlsx"
            workbook.save(workbook_path)
            yield workbook_path

    def create_temp_docx(self, build_callback):
        try:
            from docx import Document
        except ModuleNotFoundError:
            self.skipTest("python-docx is not installed")

        with tempfile.TemporaryDirectory() as temp_dir:
            document = Document()
            build_callback(document)
            document_path = Path(temp_dir) / "test.docx"
            document.save(document_path)
            yield document_path

    def create_temp_pdf(self, build_callback):
        try:
            from reportlab.pdfgen import canvas
        except ModuleNotFoundError:
            self.skipTest("reportlab is not installed")

        with tempfile.TemporaryDirectory() as temp_dir:
            pdf_path = Path(temp_dir) / "test.pdf"
            pdf_canvas = canvas.Canvas(str(pdf_path), pagesize=(612, 792))
            build_callback(pdf_canvas)
            pdf_canvas.save()
            yield pdf_path

    def test_parse_splitter_source_builds_markdown_elements(self) -> None:
        markdown_text = """# 文档标题

## 第一节

第一节正文。

- 列表项
"""

        source = parse_splitter_source(markdown_text, "md")

        self.assertEqual(source.file_type, "md")
        self.assertIsNotNone(source.elements)
        self.assertEqual(
            [element.element_type for element in source.elements or []],
            ["heading", "heading", "paragraph", "list"],
        )
        self.assertEqual((source.elements or [])[1].metadata["heading_path"], ["文档标题", "第一节"])
        self.assertEqual((source.elements or [])[2].metadata["source_parser"], "markdown_parser")

    def test_parse_splitter_source_builds_plain_text_elements(self) -> None:
        text = """第一章 总则

这里是第一章内容。

第二章 范围

这里是第二章内容。
"""

        source = parse_splitter_source(text, "txt")

        self.assertEqual(source.file_type, "txt")
        self.assertIsNotNone(source.elements)
        self.assertEqual(
            [element.element_type for element in source.elements or []],
            ["heading", "paragraph", "heading", "paragraph"],
        )
        self.assertEqual((source.elements or [])[0].level, 1)
        self.assertEqual((source.elements or [])[1].metadata["heading_path"], ["第一章 总则"])
        self.assertEqual((source.elements or [])[2].metadata["source_parser"], "plain_text_parser")

    def test_parse_splitter_source_builds_csv_table_element(self) -> None:
        csv_text = """id,name,score
1,Alice,95
2,Bob,88
"""

        source = parse_splitter_source(csv_text, "csv")

        self.assertEqual(source.file_type, "csv")
        self.assertIsNotNone(source.elements)
        self.assertEqual(len(source.elements or []), 1)

        element = (source.elements or [])[0]
        self.assertEqual(element.element_type, "table")
        self.assertEqual(element.metadata["source_parser"], "csv_parser")
        self.assertTrue(element.metadata["has_header"])
        self.assertEqual(element.metadata["row_start"], 2)
        self.assertEqual(element.metadata["row_end"], 3)
        self.assertEqual(element.metadata["col_start"], "A")
        self.assertEqual(element.metadata["col_end"], "C")
        self.assertIn("| id | name | score |", element.text)

    def test_parse_splitter_source_builds_excel_table_elements(self) -> None:
        def build_workbook(workbook):
            sheet1 = workbook.active
            sheet1.title = "Orders"
            sheet1.append(["id", "customer", "amount"])
            sheet1.append([1, "Alice", 95])
            sheet1.append([2, "Bob", 88])
            sheet1.append([None, None, None])
            sheet1.append(["product", "stock", "price"])
            sheet1.append(["Keyboard", 30, 199])

            sheet2 = workbook.create_sheet("Summary")
            sheet2["B2"] = "metric"
            sheet2["C2"] = "value"
            sheet2["B3"] = "total_orders"
            sheet2["C3"] = 2

        for workbook_path in self.create_temp_workbook(build_workbook):
            source = parse_splitter_source("", "xlsx", spreadsheet_path=str(workbook_path))

        self.assertEqual(source.file_type, "xlsx")
        self.assertIsNotNone(source.elements)
        self.assertEqual(len(source.elements or []), 3)

        first_element = (source.elements or [])[0]
        self.assertEqual(first_element.metadata["source_parser"], "excel_parser")
        self.assertEqual(first_element.metadata["sheet_name"], "Orders")
        self.assertEqual(first_element.metadata["row_start"], 2)
        self.assertEqual(first_element.metadata["row_end"], 3)
        self.assertEqual(first_element.metadata["col_start"], "A")
        self.assertEqual(first_element.metadata["col_end"], "C")

        last_element = (source.elements or [])[-1]
        self.assertEqual(last_element.metadata["sheet_name"], "Summary")
        self.assertEqual(last_element.metadata["sheet_used_range"], "B2:C3")

    def test_parse_splitter_source_builds_docx_elements(self) -> None:
        def build_docx(document):
            document.add_heading("员工手册", level=1)
            document.add_paragraph("第一段正文。")

            list_style = None
            for style_name in ("List Bullet", "列表项目符号", "List Paragraph"):
                try:
                    list_style = document.styles[style_name]
                    break
                except KeyError:
                    continue
            if list_style is not None:
                document.add_paragraph("准备材料", style=list_style)
                document.add_paragraph("提交审批", style=list_style)
            else:
                document.add_paragraph("准备材料")
                document.add_paragraph("提交审批")

            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "字段"
            table.cell(0, 1).text = "说明"
            table.cell(1, 0).text = "状态"
            table.cell(1, 1).text = "草稿"

        for document_path in self.create_temp_docx(build_docx):
            source = parse_splitter_source("", "docx", word_path=str(document_path))

        self.assertEqual(source.file_type, "docx")
        self.assertIsNotNone(source.elements)
        self.assertEqual(
            [element.element_type for element in source.elements or []],
            ["heading", "paragraph", "list", "table"],
        )

        heading_element = (source.elements or [])[0]
        self.assertEqual(heading_element.metadata["source_parser"], "docx_parser")
        self.assertEqual(heading_element.metadata["heading_path"], ["员工手册"])

        list_element = (source.elements or [])[2]
        self.assertIn("- 准备材料", list_element.text)
        self.assertIn("- 提交审批", list_element.text)

        table_element = (source.elements or [])[3]
        self.assertTrue(table_element.metadata["has_header"])
        self.assertEqual(table_element.metadata["col_end"], "B")
        self.assertIn("| 字段 | 说明 |", table_element.text)

    def test_parse_splitter_source_builds_pdf_layout_elements(self) -> None:
        def build_pdf(pdf_canvas):
            def draw_header_footer(page_no):
                pdf_canvas.setFont("Helvetica", 10)
                pdf_canvas.drawString(72, 770, "Internal Document")
                pdf_canvas.drawString(72, 24, f"Page {page_no}")

            draw_header_footer(1)
            pdf_canvas.setFont("Helvetica-Bold", 16)
            pdf_canvas.drawString(72, 720, "1. Overview")
            pdf_canvas.setFont("Helvetica", 12)
            pdf_canvas.drawString(72, 690, "Left column page one.")
            pdf_canvas.drawString(320, 690, "Right column page one.")
            pdf_canvas.showPage()

            draw_header_footer(2)
            pdf_canvas.setFont("Helvetica", 12)
            pdf_canvas.drawString(72, 720, "Left column page two line one.")
            pdf_canvas.drawString(72, 700, "Left column page two line two.")
            pdf_canvas.drawString(320, 720, "Right column page two line one.")
            pdf_canvas.drawString(320, 700, "Right column page two line two.")

            table_x = 72
            table_y_top = 620
            col_widths = [140, 140]
            row_height = 24
            rows = [
                ["field", "value"],
                ["status", "draft"],
                ["owner", "employee"],
            ]
            total_width = sum(col_widths)
            total_height = row_height * len(rows)
            for row_index in range(len(rows) + 1):
                y = table_y_top - row_index * row_height
                pdf_canvas.line(table_x, y, table_x + total_width, y)
            current_x = table_x
            pdf_canvas.line(current_x, table_y_top, current_x, table_y_top - total_height)
            for width in col_widths:
                current_x += width
                pdf_canvas.line(current_x, table_y_top, current_x, table_y_top - total_height)
            pdf_canvas.setFont("Helvetica", 11)
            for row_index, row in enumerate(rows):
                text_y = table_y_top - row_height * (row_index + 0.7)
                cell_x = table_x + 8
                for cell_index, cell_text in enumerate(row):
                    pdf_canvas.drawString(cell_x, text_y, cell_text)
                    cell_x += col_widths[cell_index]

        for pdf_path in self.create_temp_pdf(build_pdf):
            source = parse_splitter_source("", "pdf", pdf_path=str(pdf_path))

        self.assertEqual(source.file_type, "pdf")
        self.assertIsNotNone(source.elements)
        element_types = [element.element_type for element in source.elements or []]
        self.assertIn("heading", element_types)
        self.assertIn("paragraph", element_types)
        self.assertIn("table", element_types)
        self.assertFalse(any("Internal Document" in (element.text or "") for element in source.elements or []))
        self.assertFalse(any((element.text or "").startswith("Page ") for element in source.elements or []))
        table_elements = [element for element in source.elements or [] if element.element_type == "table"]
        self.assertTrue(table_elements)
        self.assertIn("| field | value |", table_elements[0].text)
        self.assertIsNotNone(table_elements[0].bbox)

    def test_excel_sections_split_by_sheet_context(self) -> None:
        def build_workbook(workbook):
            sheet1 = workbook.active
            sheet1.title = "Orders"
            sheet1.append(["id", "customer"])
            sheet1.append([1, "Alice"])

            sheet2 = workbook.create_sheet("Summary")
            sheet2.append(["metric", "value"])
            sheet2.append(["total_orders", 1])

        for workbook_path in self.create_temp_workbook(build_workbook):
            source = parse_splitter_source("", "xlsx", spreadsheet_path=str(workbook_path))
            sections = build_document_sections(source)
            chunks = split_document_text_pipeline("", "xlsx", spreadsheet_path=str(workbook_path))

        self.assertEqual(len(sections), 2)
        self.assertEqual([section.heading_path for section in sections], [["Orders"], ["Summary"]])
        self.assertEqual(sections[0].metadata["splitter"], "xlsx_context_structure")
        self.assertTrue(chunks[0].content.startswith("# Orders"))
        self.assertTrue(chunks[1].content.startswith("# Summary"))

    def test_docx_and_pdf_chunks_use_heading_prefix(self) -> None:
        def build_docx(document):
            document.add_heading("员工手册", level=1)
            document.add_paragraph("第一段正文。")

        for document_path in self.create_temp_docx(build_docx):
            docx_chunks = split_document_text_pipeline("", "docx", word_path=str(document_path))

        self.assertTrue(docx_chunks[0].content.startswith("# 员工手册"))

        def build_pdf(pdf_canvas):
            pdf_canvas.setFont("Helvetica", 10)
            pdf_canvas.drawString(72, 24, "Page 1")
            pdf_canvas.setFont("Helvetica-Bold", 16)
            pdf_canvas.drawString(72, 720, "1. Overview")
            pdf_canvas.setFont("Helvetica", 12)
            pdf_canvas.drawString(72, 690, "Review rules summary.")

        for pdf_path in self.create_temp_pdf(build_pdf):
            pdf_chunks = split_document_text_pipeline("", "pdf", pdf_path=str(pdf_path))

        self.assertTrue(pdf_chunks[0].content.startswith("# 1. Overview"))
        self.assertFalse(any("Page 1" in chunk.content for chunk in pdf_chunks))

    def test_markdown_sections_and_blocks_keep_structure_boundaries(self) -> None:
        markdown_text = """# 第一章

第一段内容。

- 列表 1
- 列表 2
  列表续行

| 列 1 | 列 2 |
| --- | --- |
| A | B |

```python
print("hello")
print("world")
```

## 第二节

第二节正文。
"""

        sections = split_markdown_sections(markdown_text)

        self.assertEqual(len(sections), 2)
        self.assertEqual(sections[0].heading_path, ["第一章"])
        self.assertEqual(sections[1].heading_path, ["第一章", "第二节"])

        first_section_block_types = [block.block_type for block in sections[0].blocks]
        self.assertEqual(
            first_section_block_types,
            ["heading", "paragraph", "list", "table", "code"],
        )

        self.assertEqual(
            sections[0].blocks[3].content,
            "| 列 1 | 列 2 |\n| --- | --- |\n| A | B |",
        )
        self.assertEqual(
            sections[0].blocks[4].content,
            '```python\nprint("hello")\nprint("world")\n```',
        )

        flattened_blocks = flatten_sections_to_blocks(sections)
        self.assertEqual(flattened_blocks[0].metadata["heading_path"], ["第一章"])
        self.assertEqual(flattened_blocks[-1].metadata["heading_path"], ["第一章", "第二节"])

    def test_markdown_uses_h2_as_default_section_boundary(self) -> None:
        markdown_text = """# 文档标题

总览内容。

## 第一节

第一节正文。

### 第一节-子节A

A内容。

### 第一节-子节B

B内容。

## 第二节

第二节正文。
"""

        sections = split_markdown_sections(markdown_text)

        self.assertEqual(len(sections), 3)
        self.assertEqual(sections[0].heading_path, ["文档标题"])
        self.assertEqual(sections[1].heading_path, ["文档标题", "第一节"])
        self.assertEqual(sections[2].heading_path, ["文档标题", "第二节"])

        flattened_blocks = flatten_sections_to_blocks(sections)
        subsection_paragraphs = [
            block
            for block in flattened_blocks
            if block.block_type == "paragraph" and block.content in {"A内容。", "B内容。"}
        ]
        self.assertEqual(
            [block.metadata["section_index"] for block in subsection_paragraphs],
            [1, 1],
        )
        self.assertEqual(
            [block.metadata["heading_path"] for block in subsection_paragraphs],
            [
                ["文档标题", "第一节", "第一节-子节A"],
                ["文档标题", "第一节", "第一节-子节B"],
            ],
        )

    def test_markdown_falls_back_to_h1_when_no_h2_exists(self) -> None:
        markdown_text = """# 第一章

第一章内容。

# 第二章

第二章内容。
"""

        sections = split_markdown_sections(markdown_text)

        self.assertEqual(len(sections), 2)
        self.assertEqual(sections[0].heading_path, ["第一章"])
        self.assertEqual(sections[1].heading_path, ["第二章"])

    def test_heading_context_before_first_h2_does_not_create_section(self) -> None:
        markdown_text = """# 文档标题

## 第一节

第一节正文。

## 第二节

第二节正文。
"""

        sections = split_markdown_sections(markdown_text)

        self.assertEqual(len(sections), 2)
        self.assertEqual(sections[0].heading_path, ["文档标题", "第一节"])
        self.assertEqual(sections[1].heading_path, ["文档标题", "第二节"])

    def test_heading_context_with_preface_body_still_creates_section(self) -> None:
        markdown_text = """# 文档标题

导语内容。

## 第一节

第一节正文。
"""

        sections = split_markdown_sections(markdown_text)

        self.assertEqual(len(sections), 2)
        self.assertEqual(sections[0].heading_path, ["文档标题"])
        self.assertEqual(sections[1].heading_path, ["文档标题", "第一节"])

    def test_is_only_heading_context_section(self) -> None:
        self.assertTrue(
            is_only_heading_context_section(
                ["# 文档标题", "", "   "],
                section_boundary_level=2,
            )
        )
        self.assertFalse(
            is_only_heading_context_section(
                ["# 文档标题", "", "导语内容。"],
                section_boundary_level=2,
            )
        )
        self.assertFalse(
            is_only_heading_context_section(
                ["## 第一节"],
                section_boundary_level=2,
            )
        )

    def test_plain_text_becomes_paragraph_blocks(self) -> None:
        text = "第一段第一行\n第一段第二行\n\n第二段\n\n第三段"

        sections = split_plain_text_sections(text, "txt")

        self.assertEqual(len(sections), 1)
        self.assertEqual(sections[0].heading_path, [])
        self.assertEqual(
            [block.block_type for block in sections[0].blocks],
            ["paragraph", "paragraph", "paragraph"],
        )
        self.assertEqual(sections[0].blocks[0].content, "第一段第一行\n第一段第二行")

    def test_plain_text_headings_create_multiple_sections(self) -> None:
        text = """第一章 总则

这里是第一章内容。

第二章 范围

这里是第二章内容。
"""

        sections = split_plain_text_sections(text, "txt")

        self.assertEqual(len(sections), 2)
        self.assertEqual(sections[0].heading_path, ["第一章 总则"])
        self.assertEqual(sections[1].heading_path, ["第二章 范围"])
        self.assertEqual(sections[0].blocks[0].block_type, "heading")
        self.assertEqual(sections[0].blocks[1].content, "这里是第一章内容。")

    def test_plain_text_without_reliable_multiple_headings_falls_back(self) -> None:
        text = """1. 公司应当建立制度。
2. 公司应当持续改进。

这是后续说明段落。
"""

        sections = split_plain_text_sections(text, "txt")

        self.assertEqual(len(sections), 1)
        self.assertEqual(sections[0].heading_path, [])
        self.assertTrue(all(block.block_type == "paragraph" for block in sections[0].blocks))

    def test_detect_plain_text_headings(self) -> None:
        lines = [
            "第一章 总则",
            "",
            "这里是第一章内容。",
            "",
            "1.1 适用范围",
            "",
            "这里是适用范围内容。",
        ]

        self.assertEqual(
            detect_plain_text_headings(lines),
            [
                (0, 1, "第一章 总则"),
                (4, 2, "1.1 适用范围"),
            ],
        )

    def test_detect_plain_text_headings_tolerates_ocr_spacing(self) -> None:
        lines = [
            "第 一 章 总 则",
            "",
            "这里是第一章内容。",
            "",
            "1 . 1 适 用 范 围",
            "",
            "这里是适用范围内容。",
        ]

        self.assertEqual(
            detect_plain_text_headings(lines),
            [
                (0, 1, "第一章 总则"),
                (4, 2, "1.1 适用范围"),
            ],
        )

    def test_detect_plain_text_headings_tolerates_missing_blank_lines(self) -> None:
        lines = [
            "第一章 总则",
            "这里是第一章内容。",
            "",
            "第二章 范围",
            "这里是第二章内容。",
        ]

        self.assertEqual(
            detect_plain_text_headings(lines),
            [
                (0, 1, "第一章 总则"),
                (3, 1, "第二章 范围"),
            ],
        )

    def test_plain_text_cn_enum_headings_do_not_form_false_parent_child_chain(self) -> None:
        text = """一、适用范围

这里是适用范围内容。

二、审批要求

这里是审批要求内容。
"""

        source = parse_splitter_source(text, "txt")
        elements = source.elements or []
        heading_elements = [element for element in elements if element.element_type == "heading"]

        self.assertEqual(
            [element.metadata.get("heading_path") for element in heading_elements],
            [["一、适用范围"], ["二、审批要求"]],
        )

    def test_plain_text_outline_without_body_falls_back_to_paragraphs(self) -> None:
        text = """第一章 总则

第二章 范围

第三章 术语
"""

        sections = split_plain_text_sections(text, "txt")

        self.assertEqual(len(sections), 1)
        self.assertEqual(sections[0].heading_path, [])
        self.assertEqual(
            [block.content for block in sections[0].blocks],
            ["第一章 总则", "第二章 范围", "第三章 术语"],
        )

    def test_pdf_becomes_page_sections_with_paragraph_blocks(self) -> None:
        pages = [
            PdfPageText(page_number=1, text="第一页第一段\n\n第一页第二段"),
            PdfPageText(page_number=2, text="第二页唯一段"),
        ]

        sections = split_pdf_sections(pages)

        self.assertEqual(len(sections), 2)
        self.assertEqual(sections[0].metadata["page_start"], 1)
        self.assertEqual(sections[1].metadata["page_start"], 2)
        self.assertEqual(
            [block.block_type for block in sections[0].blocks],
            ["paragraph", "paragraph"],
        )
        self.assertEqual(sections[0].blocks[0].metadata["page_start"], 1)
        self.assertEqual(sections[1].blocks[0].content, "第二页唯一段")

    def test_pdf_uses_plain_text_heading_detection_when_reliable(self) -> None:
        pages = [
            PdfPageText(page_number=1, text="第一章 总则\n这里是第一章内容。"),
            PdfPageText(page_number=2, text="第二章 范围\n这里是第二章内容。"),
        ]

        sections = split_pdf_sections(pages)

        self.assertEqual(len(sections), 2)
        self.assertEqual(sections[0].heading_path, ["第一章 总则"])
        self.assertEqual(sections[1].heading_path, ["第二章 范围"])
        self.assertEqual(sections[0].blocks[0].block_type, "heading")
        self.assertEqual(sections[0].blocks[1].metadata["page_start"], 1)
        self.assertEqual(sections[1].blocks[1].metadata["page_start"], 2)

    def test_pdf_single_heading_builds_cross_page_section(self) -> None:
        pages = [
            PdfPageText(page_number=1, text="第一章 总则\n第一页正文。"),
            PdfPageText(page_number=2, text="第二页正文。"),
            PdfPageText(page_number=3, text="第三页正文。"),
        ]

        sections = split_pdf_sections(pages)

        self.assertEqual(len(sections), 1)
        self.assertEqual(sections[0].heading_path, ["第一章 总则"])
        self.assertEqual(sections[0].metadata["page_start"], 1)
        self.assertEqual(sections[0].metadata["page_end"], 3)
        self.assertEqual(
            [block.metadata.get("page_start") for block in sections[0].blocks],
            [1, 1, 2, 3],
        )

    def test_pdf_without_headings_keeps_page_fallback(self) -> None:
        pages = [
            PdfPageText(page_number=1, text="第一页第一段\n\n第一页第二段"),
            PdfPageText(page_number=2, text="第二页唯一段"),
            PdfPageText(page_number=3, text="第三页唯一段"),
        ]

        sections = split_pdf_sections(pages)

        self.assertEqual(len(sections), 3)
        self.assertEqual([section.metadata["page_start"] for section in sections], [1, 2, 3])
        self.assertEqual([section.metadata["page_end"] for section in sections], [1, 2, 3])

    def test_pdf_layout_chunks_keep_table_and_strip_repeated_headers(self) -> None:
        def build_pdf(pdf_canvas):
            def draw_header_footer(page_no):
                pdf_canvas.setFont("Helvetica", 10)
                pdf_canvas.drawString(72, 770, "Internal Document")
                pdf_canvas.drawString(72, 24, f"Page {page_no}")

            draw_header_footer(1)
            pdf_canvas.setFont("Helvetica-Bold", 16)
            pdf_canvas.drawString(72, 720, "1. Overview")
            pdf_canvas.setFont("Helvetica", 12)
            pdf_canvas.drawString(72, 690, "Left column page one.")
            pdf_canvas.drawString(320, 690, "Right column page one.")
            pdf_canvas.showPage()

            draw_header_footer(2)
            pdf_canvas.setFont("Helvetica", 12)
            pdf_canvas.drawString(72, 720, "Left column page two line one.")
            pdf_canvas.drawString(72, 700, "Left column page two line two.")
            pdf_canvas.drawString(320, 720, "Right column page two line one.")
            pdf_canvas.drawString(320, 700, "Right column page two line two.")

            table_x = 72
            table_y_top = 620
            col_widths = [140, 140]
            row_height = 24
            rows = [
                ["field", "value"],
                ["status", "draft"],
                ["owner", "employee"],
                ["reviewer", "manager"],
            ]
            total_width = sum(col_widths)
            total_height = row_height * len(rows)
            for row_index in range(len(rows) + 1):
                y = table_y_top - row_index * row_height
                pdf_canvas.line(table_x, y, table_x + total_width, y)
            current_x = table_x
            pdf_canvas.line(current_x, table_y_top, current_x, table_y_top - total_height)
            for width in col_widths:
                current_x += width
                pdf_canvas.line(current_x, table_y_top, current_x, table_y_top - total_height)
            pdf_canvas.setFont("Helvetica", 11)
            for row_index, row in enumerate(rows):
                text_y = table_y_top - row_height * (row_index + 0.7)
                cell_x = table_x + 8
                for cell_index, cell_text in enumerate(row):
                    pdf_canvas.drawString(cell_x, text_y, cell_text)
                    cell_x += col_widths[cell_index]

        for pdf_path in self.create_temp_pdf(build_pdf):
            chunks = split_document_text_pipeline(
                "",
                "pdf",
                target_chunk_size=80,
                max_chunk_size=120,
                chunk_overlap=20,
                pdf_path=str(pdf_path),
            )

        self.assertGreaterEqual(len(chunks), 2)
        self.assertFalse(any("Internal Document" in chunk.content for chunk in chunks))
        self.assertTrue(any(chunk.metadata.get("block_type") == "table" for chunk in chunks))
        table_chunks = [chunk for chunk in chunks if chunk.metadata.get("block_type") == "table"]
        self.assertTrue(all("| field | value |" in chunk.content for chunk in table_chunks))
        self.assertTrue(any(chunk.metadata.get("page_start") == 2 for chunk in table_chunks))

    def test_pdf_layout_does_not_misclassify_single_column_long_lines_as_two_columns(self) -> None:
        def build_pdf(pdf_canvas):
            pdf_canvas.setFont("Helvetica-Bold", 16)
            pdf_canvas.drawString(72, 740, "Procurement Supplier Policy")

            pdf_canvas.setFont("Helvetica", 11)
            pdf_canvas.drawString(
                72,
                710,
                "Purpose and scope this policy defines procurement intake supplier review contract approval delivery checks and payment controls across internal teams and external vendors.",
            )
            pdf_canvas.drawString(
                72,
                690,
                "Workflow summary business owners submit a request procurement compares options legal reviews terms security validates access design and finance confirms payment rules before approval.",
            )
            pdf_canvas.drawString(
                72,
                670,
                "Check items summary 1. business license and delivery cases are complete.",
            )
            pdf_canvas.drawString(
                72,
                650,
                "2. access control logging and vulnerability SLA pass security review.",
            )
            pdf_canvas.drawString(
                72,
                630,
                "3. confidentiality clauses and subcontracting limits are explicit in contract.",
            )
            pdf_canvas.drawString(
                72,
                610,
                "4. on call support escalation process and project staffing cover key scenarios.",
            )
            pdf_canvas.drawString(
                72,
                590,
                "Risk note vendors that cannot explain permission management or refuse breach responsibility should be treated as high risk.",
            )

        for pdf_path in self.create_temp_pdf(build_pdf):
            source = parse_splitter_source("", "pdf", pdf_path=str(pdf_path))

        elements = source.elements or []
        heading_elements = [element for element in elements if element.element_type == "heading"]
        paragraph_elements = [element for element in elements if element.element_type == "paragraph"]

        self.assertEqual(len(heading_elements), 1)
        self.assertEqual(heading_elements[0].text, "Procurement Supplier Policy")
        self.assertTrue(paragraph_elements)
        self.assertTrue(all(element.metadata.get("column_index") == 0 for element in elements))
        self.assertTrue(all(not element.text.startswith("2.") for element in heading_elements))
        combined_text = " ".join(element.text for element in paragraph_elements)
        self.assertIn("Purpose and scope", combined_text)
        self.assertIn("Workflow summary", combined_text)
        self.assertIn("Risk note", combined_text)

    def test_pdf_layout_does_not_treat_numbered_list_items_as_headings(self) -> None:
        def build_pdf(pdf_canvas):
            pdf_canvas.setFont("Helvetica-Bold", 16)
            pdf_canvas.drawString(72, 740, "Review Checklist")
            pdf_canvas.setFont("Helvetica", 11)
            pdf_canvas.drawString(72, 710, "1. Confirm project background and owner.")
            pdf_canvas.drawString(72, 690, "2. Confirm approval chain and budget source.")
            pdf_canvas.drawString(72, 670, "3. Confirm security and legal review status.")
            pdf_canvas.drawString(72, 650, "4. Confirm rollback plan and notification list.")

        for pdf_path in self.create_temp_pdf(build_pdf):
            chunks = split_document_text_pipeline(
                "",
                "pdf",
                target_chunk_size=200,
                max_chunk_size=260,
                chunk_overlap=40,
                pdf_path=str(pdf_path),
            )

        self.assertEqual(len(chunks), 1)
        self.assertEqual(chunks[0].metadata.get("heading_path"), ["Review Checklist"])
        self.assertNotIn("# 2. Confirm approval chain", chunks[0].content)
        self.assertIn("2. Confirm approval chain and budget source.", chunks[0].content)

    def test_pdf_layout_document_to_text_preserves_reading_order(self) -> None:
        def build_pdf(pdf_canvas):
            pdf_canvas.setFont("Helvetica-Bold", 16)
            pdf_canvas.drawString(72, 740, "Procurement Supplier Policy")
            pdf_canvas.setFont("Helvetica", 11)
            pdf_canvas.drawString(
                72,
                710,
                "Purpose and scope define intake review approval and payment control.",
            )
            pdf_canvas.drawString(
                72,
                690,
                "Workflow summary business submits request procurement compares options.",
            )
            pdf_canvas.drawString(
                72,
                670,
                "Risk note vendors without permission model should be treated as high risk.",
            )

        for pdf_path in self.create_temp_pdf(build_pdf):
            text = pdf_layout_document_to_text(str(pdf_path))

        self.assertIsNotNone(text)
        self.assertIn("Procurement Supplier Policy", text)
        self.assertLess(
            text.index("Purpose and scope"),
            text.index("Workflow summary"),
        )
        self.assertLess(
            text.index("Workflow summary"),
            text.index("Risk note"),
        )

    def test_extract_text_from_file_uses_layout_aware_pdf_text_when_available(self) -> None:
        def build_pdf(pdf_canvas):
            pdf_canvas.setFont("Helvetica-Bold", 16)
            pdf_canvas.drawString(72, 740, "Review Checklist")
            pdf_canvas.setFont("Helvetica", 11)
            pdf_canvas.drawString(72, 710, "1. Confirm project background and owner.")
            pdf_canvas.drawString(72, 690, "2. Confirm approval chain and budget source.")
            pdf_canvas.drawString(72, 670, "3. Confirm security and legal review status.")

        for pdf_path in self.create_temp_pdf(build_pdf):
            text = extract_text_from_file(pdf_path, ".pdf")

        self.assertIn("Review Checklist", text)
        self.assertIn("1. Confirm project background and owner.", text)
        self.assertIn("2. Confirm approval chain and budget source.", text)
        self.assertIn("3. Confirm security and legal review status.", text)

    def test_build_pdf_paragraph_text_normalizes_visual_line_spaces(self) -> None:
        lines = [
            PdfLayoutLine(
                text="本制度用于规范采购申请、合同签署 、到货验收流程。",
                page_number=1,
                bbox=[72, 100, 300, 112],
                avg_font_size=10,
                column_index=0,
            ),
            PdfLayoutLine(
                text="IT 设 备采购和供应 商管理需要同步留痕。 2 . 权限模型应复核。",
                page_number=1,
                bbox=[72, 116, 300, 128],
                avg_font_size=10,
                column_index=0,
            ),
        ]

        text = build_pdf_paragraph_text(lines)

        self.assertIn("合同签署、到货验收流程。", text)
        self.assertIn("IT设备采购和供应商管理需要同步留痕。", text)
        self.assertIn("2. 权限模型应复核。", text)

    def test_pdf_noise_paragraph_detects_short_fragment(self) -> None:
        self.assertTrue(is_probable_pdf_noise_paragraph("、表格识别与页脚去噪。"))
        self.assertFalse(is_probable_pdf_noise_paragraph("审批记录需要保留申请人、审批链和附件。"))

    def test_markdown_chunks_keep_heading_prefix_and_do_not_cross_heading(self) -> None:
        markdown_text = """# 第一章

这是第一章第一句。这是第一章第二句。这是第一章第三句。这是第一章第四句。这是第一章第五句。这是第一章第六句。

## 第二节

这是第二节第一句。这是第二节第二句。这是第二节第三句。这是第二节第四句。这是第二节第五句。
"""

        chunks = split_document_text(
            markdown_text,
            "md",
            target_chunk_size=20,
            max_chunk_size=30,
            chunk_overlap=10,
        )

        self.assertGreaterEqual(len(chunks), 4)
        self.assertTrue(all(chunk.content.startswith("#") for chunk in chunks))

        first_section_chunks = [chunk for chunk in chunks if chunk.metadata["heading_path"] == ["第一章"]]
        second_section_chunks = [
            chunk
            for chunk in chunks
            if chunk.metadata["heading_path"] == ["第一章", "第二节"]
        ]

        self.assertTrue(first_section_chunks)
        self.assertTrue(second_section_chunks)
        self.assertTrue(all(chunk.content.startswith("# 第一章") for chunk in first_section_chunks))
        self.assertTrue(all(chunk.content.startswith("## 第二节") for chunk in second_section_chunks))
        self.assertTrue(all("## 第二节" not in chunk.content for chunk in first_section_chunks))

    def test_phase_one_pipeline_splitter_handles_markdown(self) -> None:
        markdown_text = """# 文档标题

## 第一节

第一句。第二句。第三句。
"""

        chunks = split_document_text_pipeline(
            markdown_text,
            "md",
            target_chunk_size=18,
            max_chunk_size=26,
            chunk_overlap=8,
        )

        self.assertGreaterEqual(len(chunks), 1)
        self.assertTrue(all(chunk.content.startswith("## 第一节") for chunk in chunks))

    def test_csv_chunks_preserve_header_and_row_ranges(self) -> None:
        csv_text = """id,name,score
1,Alice Wonderland,95
2,Bob Robertson,88
3,Charlie Johnson,91
4,Denise Thompson,86
5,Edward Williams,93
6,Fiona Garcia,90
"""

        chunks = split_document_text_pipeline(
            csv_text,
            "csv",
            target_chunk_size=70,
            max_chunk_size=95,
            chunk_overlap=20,
        )

        self.assertGreaterEqual(len(chunks), 2)
        self.assertTrue(all("| id | name | score |" in chunk.content for chunk in chunks))
        self.assertTrue(all("| --- | --- | --- |" in chunk.content for chunk in chunks))
        self.assertTrue(all(chunk.metadata["has_header"] for chunk in chunks))
        self.assertEqual(chunks[0].metadata["row_start"], 2)
        self.assertGreaterEqual(chunks[0].metadata["row_end"], chunks[0].metadata["row_start"])
        self.assertGreater(chunks[1].metadata["row_start"], chunks[0].metadata["row_start"])
        self.assertEqual(chunks[-1].metadata["row_end"], 7)

    def test_excel_chunks_preserve_sheet_name_header_and_row_ranges(self) -> None:
        def build_workbook(workbook):
            sheet1 = workbook.active
            sheet1.title = "Sales"
            sheet1.append(["id", "customer", "amount"])
            sheet1.append([1, "Alice Wonderland", 95])
            sheet1.append([2, "Bob Robertson", 88])
            sheet1.append([3, "Charlie Johnson", 91])
            sheet1.append([4, "Denise Thompson", 86])
            sheet1.append([5, "Edward Williams", 93])
            sheet1.append([6, "Fiona Garcia", 90])

            sheet2 = workbook.create_sheet("Inventory")
            sheet2.append(["sku", "name", "stock"])
            sheet2.append(["K-01", "Keyboard", 30])
            sheet2.append(["M-02", "Mouse", 50])

        for workbook_path in self.create_temp_workbook(build_workbook):
            chunks = split_document_text_pipeline(
                "",
                "xlsx",
                target_chunk_size=70,
                max_chunk_size=95,
                chunk_overlap=20,
                spreadsheet_path=str(workbook_path),
            )

        self.assertGreaterEqual(len(chunks), 3)
        sales_chunks = [chunk for chunk in chunks if chunk.metadata.get("sheet_name") == "Sales"]
        inventory_chunks = [chunk for chunk in chunks if chunk.metadata.get("sheet_name") == "Inventory"]
        self.assertTrue(sales_chunks)
        self.assertTrue(inventory_chunks)
        self.assertTrue(all("| id | customer | amount |" in chunk.content for chunk in sales_chunks))
        self.assertTrue(all(chunk.metadata["has_header"] for chunk in sales_chunks))
        self.assertEqual(sales_chunks[0].metadata["row_start"], 2)
        self.assertEqual(sales_chunks[-1].metadata["row_end"], 7)
        self.assertEqual(inventory_chunks[0].metadata["sheet_used_range"], "A1:C3")

    def test_docx_chunks_keep_heading_list_and_table_structure(self) -> None:
        def build_docx(document):
            document.add_heading("报销制度", level=1)
            document.add_paragraph("员工提交报销前需要完成审批。")

            list_style = None
            for style_name in ("List Bullet", "列表项目符号", "List Paragraph"):
                try:
                    list_style = document.styles[style_name]
                    break
                except KeyError:
                    continue
            if list_style is not None:
                document.add_paragraph("准备发票", style=list_style)
                document.add_paragraph("填写金额", style=list_style)
            else:
                document.add_paragraph("准备发票")
                document.add_paragraph("填写金额")

            table = document.add_table(rows=4, cols=2)
            table.cell(0, 0).text = "字段"
            table.cell(0, 1).text = "说明"
            table.cell(1, 0).text = "状态"
            table.cell(1, 1).text = "草稿"
            table.cell(2, 0).text = "提交人"
            table.cell(2, 1).text = "员工"
            table.cell(3, 0).text = "审批人"
            table.cell(3, 1).text = "主管"

        for document_path in self.create_temp_docx(build_docx):
            chunks = split_document_text_pipeline(
                "",
                "docx",
                target_chunk_size=60,
                max_chunk_size=85,
                chunk_overlap=20,
                word_path=str(document_path),
            )

        self.assertGreaterEqual(len(chunks), 2)
        self.assertEqual(chunks[0].metadata.get("heading_path"), ["报销制度"])
        self.assertIn("员工提交报销前需要完成审批。", chunks[0].content)
        self.assertTrue(any("- 准备发票" in chunk.content for chunk in chunks))

        table_chunks = [chunk for chunk in chunks if chunk.metadata.get("block_type") == "table"]
        self.assertTrue(table_chunks)
        self.assertTrue(all("| 字段 | 说明 |" in chunk.content for chunk in table_chunks))
        self.assertEqual(table_chunks[0].metadata["row_start"], 2)
        self.assertEqual(table_chunks[-1].metadata["row_end"], 4)

    def test_semantic_overlap_does_not_start_with_half_sentence(self) -> None:
        markdown_text = """# 标题

第一句内容比较长一些。第二句内容比较长一些。第三句内容比较长一些。第四句内容比较长一些。
"""

        chunks = split_document_text(
            markdown_text,
            "md",
            target_chunk_size=18,
            max_chunk_size=26,
            chunk_overlap=8,
        )

        self.assertGreaterEqual(len(chunks), 2)

        expected_body_starts = {
            "第一句内容比较长一些。",
            "第二句内容比较长一些。",
            "第三句内容比较长一些。",
            "第四句内容比较长一些。",
        }

        for chunk in chunks:
            body = chunk.content.split("\n\n", 1)[1]
            first_body_line = body.split("\n\n", 1)[0]
            self.assertIn(first_body_line, expected_body_starts)

    def test_markdown_table_chunks_do_not_start_from_middle_row(self) -> None:
        markdown_text = """# 表格节

| 列1 | 列2 |
| --- | --- |
| 第一行很长的数据A | 第一行很长的数据B |
| 第二行很长的数据A | 第二行很长的数据B |
| 第三行很长的数据A | 第三行很长的数据B |
"""

        chunks = split_document_text(
            markdown_text,
            "md",
            target_chunk_size=40,
            max_chunk_size=70,
            chunk_overlap=16,
        )

        self.assertGreaterEqual(len(chunks), 2)
        table_header = "| 列1 | 列2 |\n| --- | --- |"
        self.assertTrue(all(table_header in chunk.content for chunk in chunks))

    def test_table_boundary_forces_chunk_flush(self) -> None:
        markdown_text = """# 文档标题

## 第一节

前言说明。

| 列1 | 列2 |
| --- | --- |
| A | B |

后续说明。
"""

        chunks = split_document_text(
            markdown_text,
            "md",
            target_chunk_size=200,
            max_chunk_size=400,
            chunk_overlap=50,
        )

        self.assertEqual(len(chunks), 3)
        self.assertIn("前言说明。", chunks[0].content)
        self.assertNotIn("| 列1 | 列2 |", chunks[0].content)
        self.assertIn("| 列1 | 列2 |", chunks[1].content)
        self.assertNotIn("前言说明。", chunks[1].content)
        self.assertNotIn("后续说明。", chunks[1].content)
        self.assertIn("后续说明。", chunks[2].content)
        self.assertNotIn("| 列1 | 列2 |", chunks[2].content)

    def test_code_boundary_forces_chunk_flush(self) -> None:
        markdown_text = """# 文档标题

## 第一节

前言说明。

```python
print("hello")
print("world")
```

后续说明。
"""

        chunks = split_document_text(
            markdown_text,
            "md",
            target_chunk_size=200,
            max_chunk_size=400,
            chunk_overlap=50,
        )

        self.assertEqual(len(chunks), 3)
        self.assertIn("前言说明。", chunks[0].content)
        self.assertNotIn("```python", chunks[0].content)
        self.assertIn("```python", chunks[1].content)
        self.assertNotIn("前言说明。", chunks[1].content)
        self.assertNotIn("后续说明。", chunks[1].content)
        self.assertIn("后续说明。", chunks[2].content)
        self.assertNotIn("```python", chunks[2].content)


if __name__ == "__main__":
    unittest.main()
