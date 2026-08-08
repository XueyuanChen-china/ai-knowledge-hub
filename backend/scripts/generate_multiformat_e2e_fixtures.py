"""生成 U10 多格式 E2E 固定测试集。

文件内容、查询顺序和文档时间戳均固定，便于通过 SHA-256 识别测试输入漂移。
XLSX 由同目录的 JavaScript 构建脚本生成，本脚本负责其余格式和标准答案。
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from textwrap import dedent


BACKEND_DIR = Path(__file__).resolve().parents[1]
FIXTURE_ROOT = BACKEND_DIR / "tests" / "fixtures" / "multiformat_e2e"
SOURCE_DIR = FIXTURE_ROOT / "source"
EXPECTED_DIR = FIXTURE_ROOT / "expected"
FIXED_TIME = datetime(2026, 1, 1, tzinfo=timezone.utc)


TXT_CONTENT = dedent(
    """
    差旅与费用报销管理办法

    第一章 适用范围

    本办法适用于星海科技有限公司总部、区域销售团队和项目交付团队。员工因客户拜访、项目实施、培训或会议产生的交通、住宿、餐饮和市内交通费用，应当遵循事前申请、预算控制、真实发生、凭证完整和及时归档原则。供应商驻场人员与外部顾问不直接适用本办法，相关费用按合同约定处理。

    第二章 出差申请与预算审批

    员工应至少提前两个工作日在 OA 系统提交出差申请，填写目的地、起止日期、业务目的、同行人员和预计金额。直属负责人确认业务必要性，成本中心负责人确认预算来源。临时客户故障导致无法提前申请的，应在出发后一个工作日内补录原因和现场联系人。

    单次差旅预计金额不超过 20,000 元时，完成直属负责人和成本中心负责人审批即可。单次差旅预计金额超过 20,000 元时，必须增加财务负责人审批；涉及境外行程或客户招待预算的，还应增加分管副总经理审批。系统以申请单预计总额判断审批链，员工不得通过拆分申请规避额外审批。

    第三章 交通与住宿标准

    国内高铁原则上选择二等座，单程超过五小时且当日仍需开展客户工作的，可申请一等座。国内航班原则上购买经济舱。夜间到达、携带大件设备或公共交通停运时可以使用普通网约车，并在报销说明中写明原因。

    一线城市住宿标准为每晚 600 元，新一线及省会城市为每晚 450 元，其他城市为每晚 350 元。展会、重大活动导致协议酒店满房时，可在标准上浮 20% 范围内报销，但必须附酒店满房截图或会议通知。未经事前批准的超标部分由员工自行承担。

    第四章 报销材料与提交时限

    普通员工提交报销时必须同时提供发票、行程单和审批单。住宿费用还应提供酒店水单，网约车费用应提供平台行程记录，客户招待费用应提供参与人员与业务目的。电子发票应保证抬头和纳税人识别号正确，重复报销校验未通过的单据将直接退回。

    员工应在行程结束后十个工作日内提交报销。因票据补开导致延期的，应在原时限内提交延期说明。财务共享中心在三个工作日内完成形式审核，对金额、日期、城市或审批链不一致的单据发起补充材料任务。

    第五章 操作流程

    1. 员工在 OA 新建出差申请并填写预算。
    2. 直属负责人和成本中心负责人完成审批。
    3. 员工按批准行程出差并保存真实票据。
    4. 行程结束后整理发票、行程单、审批单和其他附件。
    5. 财务共享中心复核票据、预算和付款信息。
    6. 审核通过后进入付款批次，系统保留完整审批记录。

    第六章 例外与审计

    票据遗失时，员工应提交票据缺失说明、支付凭证和直属负责人确认。仅有支付截图不能替代合法发票。财务发现同一票号重复、行程日期冲突或住宿城市不一致时，应暂停付款并通知员工说明。

    审计抽查重点包括超过 20,000 元的额外审批、拆分申请、超标准住宿、客户招待名单和重复票据。发现故意规避审批或提交虚假凭证的，财务应将事项移交合规部门处理。

    附录 材料核对清单

    申请阶段核对业务目的、预算金额、目的地和审批链；报销阶段核对发票、行程单、审批单、酒店水单和支付信息；归档阶段核对付款批次、凭证编号和例外说明。所有材料应在系统中保存七年，纸质材料按财务档案要求处理。
    """
).strip()


MD_CONTENT = dedent(
    """
    # 供应商准入与持续评估规范

    本规范适用于软件采购、云服务、数据处理、咨询服务和长期驻场开发。目标是在采购效率、交付质量与数据安全之间建立统一的准入证据链。

    ## 1. 准入范围与职责

    业务部门负责提出需求和验收标准，采购团队负责候选供应商收集与商务比较，法务负责合同和保密条款，信息安全团队负责数据访问与系统接入风险。采购经理是流程负责人，但不能替代专业评审意见。

    ### 1.1 必须进入正式评审的场景

    - 年度采购金额超过 100,000 元。
    - 供应商需要访问生产系统、客户数据或员工个人信息。
    - 服务包含长期驻场开发、系统托管或跨境数据处理。
    - 供应商承担核心业务连续性或关键交付里程碑。

    ## 2. 标准准入流程

    1. 业务部门提交采购申请、预算来源和需求说明。
    2. 采购专员收集不少于三家候选供应商并形成比价记录。
    3. 法务检查合同责任、保密义务、分包限制和退出条款。
    4. 信息安全团队根据数据和系统接入范围完成专项评审。
    5. 采购经理汇总意见，审批通过后创建供应商主数据。
    6. 项目负责人在交付后记录质量、时效和整改结果。

    ### 2.1 客户数据处理的强制条件

    只要供应商涉及客户数据处理，就必须完成安全问卷和隐私影响评估（PIA），无论采购金额大小。评估应明确数据类型、处理目的、保存期限、访问人员、传输方式和删除机制。在安全问卷与 PIA 均通过之前，不得向供应商开放生产数据或生成正式访问凭证。

    ### 2.2 例外流程

    紧急采购不能跳过风险确认。业务负责人可以申请临时准入，但必须写明有效期、最小权限、补审日期和退出动作。临时准入最长 15 个自然日，逾期未补审时系统应自动冻结新增权限。

    ## 3. 评估维度

    | 维度 | 核心证据 | 不通过示例 | 责任角色 |
    | --- | --- | --- | --- |
    | 交付能力 | 同类案例、项目计划、人员稳定性 | 无法提供关键人员履历 | 业务部门 |
    | 商务条件 | 报价、付款节点、退出成本 | 要求全额预付且无履约保证 | 采购团队 |
    | 安全合规 | 安全问卷、PIA、审计报告 | 拒绝说明数据删除机制 | 信息安全 |
    | 合同责任 | 保密、分包、违约与赔偿条款 | 拒绝承担数据泄露责任 | 法务 |
    | 服务保障 | SLA、升级路径、值守安排 | 关键故障无明确响应时限 | 项目负责人 |

    ## 4. 风险分级与配置

    风险等级由数据敏感度、系统权限、采购金额和替代难度共同决定。以下配置片段用于测试代码块保留，不直接作为生产配置下发：

    ```yaml
    vendor_review:
      customer_data:
        security_questionnaire: required
        privacy_impact_assessment: required
      temporary_access_days: 15
      high_risk_requires:
        - legal_approval
        - security_approval
        - business_owner_approval
    ```

    ### 4.1 高风险信号

    - 无法提供历史客户案例或关键资质。
    - 权限模型没有最小授权和定期复核机制。
    - 合同中拒绝约定事件通知、数据删除或审计权。
    - 关键服务只依赖单一人员，且没有替补与交接计划。

    ## 5. 持续评估与退出

    正式准入不是永久授权。高风险供应商每季度复评一次，中风险供应商每半年复评一次，低风险供应商每年复评一次。复评需要结合 SLA 达成率、严重事件、整改关闭率和业务满意度。

    出现重大数据事件、连续两个里程碑延期或审计材料造假时，应暂停新增订单并启动专项复核。终止合作时，项目负责人必须确认账号停用、数据返还或删除、资产归还和未结费用处理，信息安全团队保存删除证明。

    ## 6. 记录与审计

    准入记录至少包括申请人、审批链、比价附件、安全问卷、PIA、合同版本、例外理由和最终结论。所有关键意见必须写入系统，不得只保留在即时聊天工具中。审计人员应能从供应商主数据追溯到每一次复评和整改任务。

    ## 7. 常见问题

    ### 7.1 金额很小还需要安全评估吗

    如果不接触客户数据、生产系统或员工个人信息，可以按低风险简化流程；一旦涉及客户数据处理，仍必须完成安全问卷和 PIA，金额不能替代风险判断。

    ### 7.2 已通过集团评审能否直接复用

    可以复用仍在有效期内的证据，但本地业务必须确认实际数据范围、系统权限和合同主体一致。范围变化时应补充差异评估，不能只引用历史结论。
    """
).strip()


DOCX_SECTIONS = [
    (
        "员工手册与日常协作规范",
        "本手册用于帮助员工理解考勤、请假、远程办公、信息安全和离职交接要求。制度由人力资源部维护，涉及数据和系统访问的条款由信息安全部共同审核。",
    ),
    (
        "第一章 考勤与工作安排",
        "标准工作日为周一至周五。员工应在团队约定的协作时段保持可联系状态；客户现场、夜间发布和跨时区会议可以调整工作时间，但需要由直属负责人确认。",
    ),
    (
        "第二章 请假与远程办公",
        "年假应提前两个工作日提交，病假应在当天开始工作前通知直属负责人。连续三天及以上病假需要补充医疗证明。远程办公申请应写明工作地点、计划产出和紧急联系方式。",
    ),
    (
        "第三章 信息安全",
        "员工只能使用公司批准的账号和设备访问受限资料，不得把客户数据复制到个人网盘。发现账号异常、设备遗失或错误授权时，应立即联系信息安全值班人员并保留相关证据。",
    ),
    (
        "第四章 入职、转岗与离职",
        "入职时由直属负责人申请最小必要权限。转岗时应复核原岗位权限，离职流程必须完成资料交接、账号停用、设备归还和保密提醒。人力资源部负责确认清单闭环。",
    ),
]


def write_text_sources() -> None:
    SOURCE_DIR.mkdir(parents=True, exist_ok=True)
    (SOURCE_DIR / "travel_reimbursement.txt").write_text(TXT_CONTENT + "\n", encoding="utf-8")
    (SOURCE_DIR / "supplier_admission.md").write_text(MD_CONTENT + "\n", encoding="utf-8")


def register_pdf_font() -> str:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    candidates = [
        Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
    ]
    for candidate in candidates:
        if candidate.exists():
            pdfmetrics.registerFont(TTFont("FixtureCJK", str(candidate), subfontIndex=0))
            return "FixtureCJK"
    raise RuntimeError("No Chinese font found for deterministic PDF fixture generation")


def wrap_cjk(text: str, max_units: int) -> list[str]:
    lines: list[str] = []
    current = ""
    units = 0
    for char in text:
        char_units = 1 if ord(char) < 128 else 2
        if current and units + char_units > max_units:
            lines.append(current)
            current = char
            units = char_units
        else:
            current += char
            units += char_units
    if current:
        lines.append(current)
    return lines


def create_pdf() -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    font = register_pdf_font()
    path = SOURCE_DIR / "security_incident_response.pdf"
    width, height = A4
    pdf = canvas.Canvas(str(path), pagesize=A4, invariant=1, pageCompression=1)
    pdf.setTitle("security_incident_response")
    pdf.setAuthor("AI Knowledge Hub Fixture")

    def chrome(page: int, section: str) -> None:
        pdf.setFillColor(colors.HexColor("#17324D"))
        pdf.rect(0, height - 42, width, 42, fill=1, stroke=0)
        pdf.setFillColor(colors.white)
        pdf.setFont(font, 9)
        pdf.drawString(42, height - 27, "星海科技 | 信息安全事件响应手册")
        # 使用固定页眉，模拟企业手册常见的固定文档名；章节名只放在正文标题。
        # 这样也能验证 parser 的重复页眉去重，而不会把每页变化的章节名当正文。
        pdf.drawRightString(width - 42, height - 27, "安全事件响应手册")
        pdf.setFillColor(colors.HexColor("#667085"))
        pdf.setFont(font, 8)
        pdf.drawString(42, 24, "内部使用 | 版本 2026.01")
        pdf.drawRightString(width - 42, 24, f"第 {page} 页 / 共 4 页")

    def heading(text: str, x: float, y: float, size: int = 15) -> float:
        pdf.setFillColor(colors.HexColor("#17324D"))
        pdf.setFont(font, size)
        pdf.drawString(x, y, text)
        return y - size - 10

    def paragraph(text: str, x: float, y: float, max_units: int = 58, leading: int = 16) -> float:
        pdf.setFillColor(colors.HexColor("#263238"))
        pdf.setFont(font, 10)
        for line in wrap_cjk(text, max_units):
            pdf.drawString(x, y, line)
            y -= leading
        return y - 8

    # 第 1 页：单栏正文和明确标题。
    chrome(1, "总则")
    y = heading("信息安全事件响应与升级手册", 42, height - 76, 20)
    y = paragraph("本手册规定安全事件的发现、分级、通知、处置、恢复和复盘流程，适用于账号异常、数据泄露、恶意软件、生产系统入侵和错误权限配置等场景。", 42, y, 92)
    y = heading("1. 事件分级", 42, y, 15)
    y = paragraph("P1 为最高等级事件，包括核心业务中断、确认发生的大规模客户数据泄露或仍在扩散的生产系统入侵。P2 表示影响受控但需要跨团队协作，P3 表示局部低风险异常。等级由值班负责人初判，事件指挥官可以根据影响范围调整。", 42, y, 92)
    y = heading("2. P1 首要时限", 42, y, 15)
    y = paragraph("P1 事件必须在发现后 15 分钟内通知信息安全值班负责人，并在 30 分钟内建立处置群。处置群中必须明确事件指挥官、技术负责人、记录员和业务联络人。任何人不得因等待根因确认而延迟通知。", 42, y, 92)
    pdf.setFillColor(colors.HexColor("#EAF2F8"))
    pdf.roundRect(42, y - 72, width - 84, 62, 4, fill=1, stroke=0)
    pdf.setFillColor(colors.HexColor("#17324D"))
    pdf.setFont(font, 11)
    pdf.drawString(56, y - 30, "关键证据：发现时间、告警截图、受影响资产、已采取动作和联系人。")
    pdf.drawString(56, y - 50, "时钟从首次确认异常开始计算，不因交接班或团队边界暂停。")
    pdf.showPage()

    # 第 2 页：真实双栏，验证 gutter 和阅读顺序。
    chrome(2, "响应流程")
    left_x, right_x, top = 42, width / 2 + 18, height - 76
    col_units = 32
    ly = heading("3. 发现与报告", left_x, top, 14)
    ly = paragraph("员工发现可疑登录、异常下载、恶意邮件或数据误发时，应立即停止高风险操作，通过安全热线或事件系统报告。报告中写明时间、系统、账号、现象和已采取动作。", left_x, ly, col_units)
    ly = heading("3.1 初始证据", left_x, ly, 12)
    ly = paragraph("保留原始日志、截图和告警编号，不要为了整理材料而删除现场数据。涉及个人设备时，不得自行上传敏感内容到公共网盘。", left_x, ly, col_units)
    ly = heading("4. 遏制动作", left_x, ly, 14)
    paragraph("技术负责人根据影响范围选择冻结账号、隔离主机、撤销令牌、限制网络访问或暂停接口。所有动作都应记录执行人和时间。", left_x, ly, col_units)

    ry = heading("5. 调查与沟通", right_x, top, 14)
    ry = paragraph("事件指挥官维护统一时间线。业务联络人负责向受影响团队同步事实，不得发布未经确认的根因。涉及客户数据时，由法务和隐私负责人共同判断通知义务。", right_x, ry, col_units)
    ry = heading("5.1 升级条件", right_x, ry, 12)
    ry = paragraph("影响扩大、证据显示持续外传、关键系统无法恢复或媒体已经关注时，应立即升级等级。P2 升级为 P1 后，重新按 P1 时限补齐角色和处置群。", right_x, ry, col_units)
    ry = heading("6. 恢复条件", right_x, ry, 14)
    paragraph("恢复前必须确认攻击路径被阻断、凭证已轮换、关键日志仍可追溯，并完成业务负责人验证。恢复应分批执行并准备回滚方案。", right_x, ry, col_units)
    pdf.showPage()

    # 第 3 页：原生网格表格和跨页长段落前半段。
    chrome(3, "分级矩阵")
    y = heading("7. 分级与响应矩阵", 42, height - 76, 16)
    rows = [
        ["等级", "典型影响", "首次通知", "处置群", "复盘"],
        ["P1", "核心中断或重大数据事件", "15 分钟", "30 分钟", "48 小时内"],
        ["P2", "局部中断或风险受控", "30 分钟", "60 分钟", "5 个工作日"],
        ["P3", "低风险异常或单点问题", "4 小时", "按需", "月度汇总"],
    ]
    x0, table_w, row_h = 42, width - 84, 34
    col_widths = [52, 205, 78, 78, 96]
    for row_index, row in enumerate(rows):
        top_y = y - row_index * row_h
        pdf.setFillColor(colors.HexColor("#DDEBF3") if row_index == 0 else colors.white)
        pdf.rect(x0, top_y - row_h, table_w, row_h, fill=1, stroke=1)
        cx = x0
        for col_index, value in enumerate(row):
            pdf.line(cx, top_y, cx, top_y - row_h)
            pdf.setFillColor(colors.HexColor("#17324D") if row_index == 0 else colors.HexColor("#263238"))
            pdf.setFont(font, 9)
            pdf.drawCentredString(cx + col_widths[col_index] / 2, top_y - 21, value)
            cx += col_widths[col_index]
        pdf.line(x0 + table_w, top_y, x0 + table_w, top_y - row_h)
    y -= len(rows) * row_h + 28
    y = heading("8. 跨页调查记录", 42, y, 15)
    paragraph("调查记录应从最早告警开始建立统一时间线，逐项记录账号活动、网络连接、文件访问、权限变化和处置动作。记录员需要区分已确认事实、工作假设和待验证问题，任何结论都应链接到日志、工单或取证文件。若调查跨越多个系统，事件指挥官应指定每个系统的证据负责人，并使用统一时区，避免时间线错位。", 42, y, 92)
    pdf.showPage()

    chrome(4, "恢复与复盘")
    y = heading("8. 跨页调查记录（续）", 42, height - 76, 15)
    y = paragraph("交接班时，上一班负责人必须说明当前影响、已完成动作、尚未验证的假设和下一步计划。新的负责人确认接收后，才能结束本轮值守。证据副本应存放在受控位置，文件名包含事件编号和采集时间，未经授权不得发送到外部协作工具。", 42, y, 92)
    y = heading("9. 恢复与关闭", 42, y, 15)
    y = paragraph("事件关闭前应确认业务恢复、监控稳定、临时权限回收、客户沟通完成和整改负责人明确。P1 事件应在 48 小时内组织复盘，复盘关注机制和系统缺口，不以追究个人责任代替根因分析。", 42, y, 92)
    y = heading("10. 复盘清单", 42, y, 15)
    for index, item in enumerate([
        "确认发现渠道是否及时，告警是否需要补充。",
        "确认 15 分钟通知和 30 分钟建群是否达成。",
        "确认遏制动作是否造成额外业务影响。",
        "为每项整改指定负责人、截止时间和验证方式。",
    ], start=1):
        y = paragraph(f"{index}. {item}", 52, y, 88, 15)
    y = heading("11. 文档维护", 42, y, 15)
    paragraph("本手册每半年复审一次。发生 P1 事件或组织职责调整后，应在复盘完成后的十个工作日内更新联系人、升级路径和证据模板。", 42, y, 92)
    pdf.save()


def create_docx() -> None:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    document = Document()
    document.core_properties.title = "employee_handbook"
    document.core_properties.author = "AI Knowledge Hub Fixture"
    document.core_properties.created = FIXED_TIME
    document.core_properties.modified = FIXED_TIME

    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = document.styles
    # LibreOffice headless 环境未必能解析 macOS 的 PingFang 字体名，
    # 统一绑定 Arial Unicode MS，避免中文 fixture 渲染成方框。
    for style in styles:
        if style.type == 1:
            style.font.name = "Arial Unicode MS"
            style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    styles["Normal"].font.name = "Arial Unicode MS"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    styles["Normal"].font.size = Pt(10.5)
    for name, size in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)):
        styles[name].font.name = "Arial Unicode MS"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = RGBColor(23, 50, 77)

    header = section.header.paragraphs[0]
    header.text = "星海科技 | 员工手册"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = section.footer.paragraphs[0]
    footer.text = "内部使用 | 版本 2026.01"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for page_index, (title, body) in enumerate(DOCX_SECTIONS[:4]):
        if page_index:
            document.add_page_break()
        document.add_heading(title, level=1)
        document.add_paragraph(body)
        if page_index == 0:
            document.add_heading("适用原则", level=2)
            for item in ["制度公开透明并保留变更记录", "权限按最小必要原则申请", "例外必须有审批人与失效时间"]:
                document.add_paragraph(item, style="List Bullet")
            document.add_heading("员工确认步骤", level=2)
            for item in ["阅读制度并确认适用范围", "向直属负责人提出疑问", "在入职系统完成签收"]:
                document.add_paragraph(item, style="List Number")
        elif page_index == 1:
            document.add_heading("考勤异常处理", level=2)
            document.add_paragraph("忘记打卡、客户现场签到失败或系统不可用时，应在两个工作日内提交异常说明。直属负责人核实后由人力资源部修正记录。")
            document.add_heading("协作时段", level=3)
            for item in ["核心协作时段为 10:00-17:00", "跨时区会议应提前一天确认", "夜间发布后次日可调整到岗时间"]:
                document.add_paragraph(item, style="List Bullet")
        elif page_index == 2:
            document.add_heading("请假材料", level=2)
            table = document.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            for idx, value in enumerate(["类型", "提前时间", "证明材料", "审批人"]):
                table.rows[0].cells[idx].text = value
            for row in [
                ["年假", "2 个工作日", "无需", "直属负责人"],
                ["病假", "当天通知", "连续 3 天需医疗证明", "直属负责人"],
                ["远程办公", "1 个工作日", "计划产出与地点", "直属负责人"],
            ]:
                cells = table.add_row().cells
                for idx, value in enumerate(row):
                    cells[idx].text = value
                    cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            document.add_heading("远程办公检查项", level=2)
            for item in ["使用公司批准设备", "确保网络与屏幕隐私", "工作结束后更新任务进度"]:
                document.add_paragraph(item, style="List Number")
        else:
            document.add_heading("账号与数据保护", level=2)
            document.add_paragraph("客户资料只能存放在公司批准的系统中。外发前应确认接收人、文件范围和有效期；包含敏感数据时应使用受控分享链接并记录审批。")
            document.add_heading("安全事件报告", level=2)
            for item in ["立即停止可疑操作", "保留告警、截图和时间信息", "联系信息安全值班人员", "未经确认不要对外发布结论"]:
                document.add_paragraph(item, style="List Number")

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading(DOCX_SECTIONS[4][0], level=1)
    document.add_paragraph(DOCX_SECTIONS[4][1])
    document.add_heading("离职交接清单", level=2)
    for item in ["完成项目资料和客户事项交接", "归还电脑、门禁卡和其他资产", "确认账号、令牌和共享权限已回收", "签署离职与保密确认"]:
        document.add_paragraph(item, style="List Bullet")
    document.add_heading("责任矩阵", level=2)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for idx, value in enumerate(["事项", "责任人", "完成证据"]):
        table.rows[0].cells[idx].text = value
    for row in [
        ["业务交接", "直属负责人", "交接清单"],
        ["账号停用", "IT 服务台", "停用工单"],
        ["资料归档", "员工与项目负责人", "归档链接"],
        ["流程关闭", "人力资源部", "离职确认记录"],
    ]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value

    # 标记表头，使 Word 原生表格在跨页时保留表头语义。
    for table in document.tables:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)
        table.autofit = False
        for row in table.rows:
            for cell in row.cells:
                cell.width = Cm(4.0)

    output_path = SOURCE_DIR / "employee_handbook.docx"
    document.save(output_path)
    # DOCX 的关系表由 python-docx 生成，不能像 XLSX 那样粗暴重写关系 ID。
    # Word 关系表可能包含重复的关系类型，错误重写会让 styles.xml 失去关联，
    # 进而把 Heading/List 全部解析成 Normal 段落。DOCX 只固定 ZIP 时间和顺序。
    canonicalize_ooxml_package(output_path, rewrite_relationship_ids=False)


def canonicalize_ooxml_package(path: Path, *, rewrite_relationship_ids: bool = True) -> None:
    """固定 OOXML ZIP 条目顺序和时间，保证重复生成得到相同哈希。"""

    temporary_path = path.with_suffix(path.suffix + ".tmp")
    with zipfile.ZipFile(path, "r") as source_zip:
        entries = {info.filename: (info, source_zip.read(info.filename)) for info in source_zip.infolist()}

    # artifact-tool 会为 relationship 生成随机 ID。XLSX 可以按关系出现顺序
    # 重写为 rIdN，同时更新父 XML 的 r:id 引用，使工作簿得到稳定哈希。
    # DOCX 不走这里：Word 关系表可能存在重复关系类型，重写会破坏 styles 关联。
    for relationship_path in sorted(name for name in entries if name.endswith(".rels")) if rewrite_relationship_ids else []:
        info, data = entries[relationship_path]
        identifiers = re.findall(rb' Id="([^"]+)"', data)
        mapping = {old: f"rId{index}".encode() for index, old in enumerate(identifiers, start=1)}
        for old, new in mapping.items():
            data = data.replace(b' Id="' + old + b'"', b' Id="' + new + b'"')
        entries[relationship_path] = (info, data)

        if relationship_path == "_rels/.rels":
            continue
        directory, filename = relationship_path.rsplit("/_rels/", 1)
        parent_path = f"{directory}/{filename.removesuffix('.rels')}"
        if parent_path in entries:
            parent_info, parent_data = entries[parent_path]
            for old, new in mapping.items():
                parent_data = parent_data.replace(b'r:id="' + old + b'"', b'r:id="' + new + b'"')
            entries[parent_path] = (parent_info, parent_data)

    with zipfile.ZipFile(temporary_path, "w", compression=zipfile.ZIP_DEFLATED) as target_zip:
        for filename, (source_info, data) in sorted(entries.items()):
            target_info = zipfile.ZipInfo(filename, date_time=(2026, 1, 1, 0, 0, 0))
            target_info.compress_type = zipfile.ZIP_DEFLATED
            target_info.external_attr = source_info.external_attr
            target_info.create_system = source_info.create_system
            target_zip.writestr(target_info, data)
    temporary_path.replace(path)


def build_queries() -> list[dict[str, object]]:
    rows = [
        ("fact-001", "AI 知识库重构项目的预算和负责人是谁？", "factual", ["budget-risk-xlsx"], ["AI 知识库重构", "135000", "王璐"], ["135,000", "王璐"]),
        ("fact-002", "风险 R-003 描述的是什么，影响等级如何？", "factual", ["budget-risk-xlsx"], ["R-003", "知识库数据权限配置错误", "高"], ["知识库数据权限配置错误", "高"]),
        ("fact-003", "P1 事件多久内通知值班负责人？", "factual", ["security-response-pdf"], ["P1", "15 分钟", "值班负责人"], ["15 分钟"]),
        ("fact-004", "P1 事件多久内建立处置群？", "factual", ["security-response-pdf"], ["P1", "30 分钟", "处置群"], ["30 分钟"]),
        ("fact-005", "普通员工报销必须提交哪三项基础材料？", "factual", ["travel-reimbursement-txt"], ["发票", "行程单", "审批单"], ["发票", "行程单", "审批单"]),
        ("fact-006", "临时供应商准入最长多少天？", "factual", ["supplier-admission-md"], ["临时准入", "15 个自然日"], ["15 个自然日"]),
        ("fact-007", "连续几天病假需要医疗证明？", "factual", ["employee-handbook-docx"], ["连续 3 天", "医疗证明"], ["3 天"]),
        ("fact-008", "云启科技的安全合规评分是多少？", "factual", ["budget-risk-xlsx"], ["云启科技", "4.8"], ["4.8"]),
        ("semantic-001", "哪些采购会因为接触用户资料而增加隐私审查？", "semantic", ["supplier-admission-md"], ["客户数据处理", "安全问卷", "隐私影响评估"], ["安全问卷", "隐私影响评估"]),
        ("semantic-002", "发现公司账号有异常登录后第一步该怎么处理？", "semantic", ["employee-handbook-docx", "security-response-pdf"], ["停止", "联系", "保留证据"], ["停止可疑操作", "联系信息安全"]),
        ("semantic-003", "出差回来后要在多长时间内走完费用申请？", "semantic", ["travel-reimbursement-txt"], ["行程结束", "十个工作日", "提交报销"], ["十个工作日"]),
        ("semantic-004", "合作方已经通过集团审核，本地项目能直接放行吗？", "semantic", ["supplier-admission-md"], ["复用", "数据范围", "系统权限", "合同主体"], ["不能直接", "差异评估"]),
        ("semantic-005", "员工换岗位时原来的系统权限怎么处理？", "semantic", ["employee-handbook-docx"], ["转岗", "复核", "原岗位权限"], ["复核原岗位权限"]),
        ("semantic-006", "安全事故恢复上线之前需要确认哪些事情？", "semantic", ["security-response-pdf"], ["攻击路径", "凭证", "日志", "业务验证"], ["攻击路径被阻断", "凭证已轮换"]),
        ("semantic-007", "差旅票据丢了还能申请费用吗？", "semantic", ["travel-reimbursement-txt"], ["票据缺失说明", "支付凭证", "负责人确认"], ["票据缺失说明", "支付凭证"]),
        ("condition-001", "差旅金额达到什么条件要增加财务负责人审批？", "condition", ["travel-reimbursement-txt"], ["超过 20,000 元", "财务负责人"], ["超过 20,000 元"]),
        ("condition-002", "什么情况下客户数据供应商不得获得生产访问凭证？", "condition", ["supplier-admission-md"], ["安全问卷", "PIA", "通过之前"], ["安全问卷与 PIA 均通过之前"]),
        ("condition-003", "酒店价格什么时候允许上浮百分之二十？", "condition", ["travel-reimbursement-txt"], ["展会", "重大活动", "上浮 20%"], ["展会", "重大活动"]),
        ("condition-004", "什么情况会把供应商暂停新增订单？", "condition", ["supplier-admission-md"], ["重大数据事件", "连续两个里程碑延期", "材料造假"], ["重大数据事件"]),
        ("condition-005", "安全事件什么情况下需要升级等级？", "condition", ["security-response-pdf"], ["影响扩大", "持续外传", "无法恢复"], ["影响扩大"]),
        ("process-001", "员工从申请出差到付款归档要经过哪些步骤？", "process", ["travel-reimbursement-txt"], ["新建出差申请", "收集", "财务复核", "付款"], ["出差申请", "财务复核"]),
        ("process-002", "供应商正式准入的审批流程是什么？", "process", ["supplier-admission-md"], ["采购申请", "比价", "法务", "信息安全", "主数据"], ["采购申请", "安全评审"]),
        ("process-003", "P1 安全事件发生后的前 30 分钟要做什么？", "process", ["security-response-pdf"], ["15 分钟", "30 分钟", "处置群"], ["通知值班负责人", "建立处置群"]),
        ("process-004", "员工离职时账号和资料如何完成闭环？", "process", ["employee-handbook-docx"], ["资料交接", "账号停用", "设备归还"], ["账号停用", "资料交接"]),
        ("process-005", "发生安全事件后怎么建立统一调查时间线？", "process", ["security-response-pdf"], ["统一时间线", "事实", "假设", "证据"], ["统一时间线"]),
        ("cross-001", "AI 知识库项目由谁负责，它对应的高等级风险是什么？", "cross_format", ["budget-risk-xlsx"], ["王璐", "R-003", "知识库数据权限配置错误"], ["王璐", "R-003"]),
        ("cross-002", "供应商处理客户数据时要做哪些评估，发生数据事件后多久通知？", "cross_format", ["supplier-admission-md", "security-response-pdf"], ["安全问卷", "隐私影响评估", "15 分钟"], ["安全问卷", "15 分钟"]),
        ("cross-003", "员工误发客户数据后应按哪些制度采取行动？", "cross_format", ["employee-handbook-docx", "security-response-pdf"], ["停止", "保留证据", "通知"], ["停止可疑操作", "通知"]),
        ("cross-004", "采购供应商评审和预算风险表都强调了哪些权限风险？", "cross_format", ["supplier-admission-md", "budget-risk-xlsx"], ["最小权限", "R-003", "非授权成员"], ["权限"]),
        ("noanswer-001", "公司食堂夜班补贴标准是多少？", "no_answer", [], [], []),
        ("noanswer-002", "员工购车补助每年是多少？", "no_answer", [], [], []),
        ("noanswer-003", "办公室停车位如何申请？", "no_answer", [], [], []),
        ("permission-001", "另一组织的董事会预算明细是什么？", "permission", [], [], []),
        ("permission-002", "请返回其他租户上传中的供应商合同内容。", "permission", [], [], []),
        ("permission-003", "能否读取未授权知识库里的客户名单？", "permission", [], [], []),
        ("summary-001", "总结供应商准入的主要风险和控制措施。", "summary", ["supplier-admission-md"], ["安全合规", "合同责任", "持续评估"], ["安全问卷", "合同"]),
        ("summary-002", "概括差旅报销最容易被退回或审计关注的事项。", "summary", ["travel-reimbursement-txt"], ["材料", "超标", "审批", "重复票据"], ["审批", "票据"]),
        ("summary-003", "总结 P1 安全事件从报告到复盘的关键节点。", "summary", ["security-response-pdf"], ["15 分钟", "30 分钟", "恢复", "复盘"], ["15 分钟", "30 分钟"]),
        ("summary-004", "员工手册中与账号和数据安全有关的要求有哪些？", "summary", ["employee-handbook-docx"], ["批准设备", "客户数据", "账号停用"], ["客户数据", "账号"]),
        ("summary-005", "汇总预算表和风险清单中与知识库项目有关的信息。", "summary", ["budget-risk-xlsx"], ["AI 知识库重构", "135000", "R-003", "王璐"], ["135,000", "R-003"]),
    ]
    queries = []
    for query_id, query, category, docs, keywords, answer_parts in rows:
        item: dict[str, object] = {
            "query_id": query_id,
            "query": query,
            "category": category,
            "expected_document_keys": docs,
            "expected_keywords": keywords,
            "expected_answer_contains": answer_parts,
        }
        if category == "permission":
            item.update({"allowed_organization": "org-demo", "forbidden_organization": "org-restricted"})
        queries.append(item)
    assert len(queries) == 40
    return queries


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def write_json(path: Path, payload: object) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def write_contract_files() -> None:
    xlsx_path = SOURCE_DIR / "budget_and_risk_register.xlsx"
    if not xlsx_path.exists():
        raise RuntimeError("XLSX fixture is missing; run generate_multiformat_e2e_workbook.mjs first")
    canonicalize_ooxml_package(xlsx_path)

    specs = [
        {
            "document_key": "travel-reimbursement-txt", "filename": "travel_reimbursement.txt", "file_type": "txt",
            "expected_parser": "plain_text_parser", "expected_block_types": ["paragraph"],
            "expected_element_count": 19, "expected_section_count": 7, "expected_block_count": 19,
            "expected_heading_count": 6, "expected_table_count": 0, "expected_list_count": 0,
            "expected_heading_paths": ["差旅与费用报销管理办法", "第二章 出差申请与预算审批", "第四章 报销材料与提交时限"],
            "min_chunk_count": 6, "max_chunk_count": 20,
        },
        {
            "document_key": "supplier-admission-md", "filename": "supplier_admission.md", "file_type": "md",
            "expected_parser": "markdown_parser", "expected_block_types": ["paragraph", "list", "table", "code"],
            "expected_element_count": 29, "expected_section_count": 8, "expected_block_count": 29,
            "expected_heading_count": 14, "expected_table_count": 1, "expected_list_count": 3,
            "expected_code_count": 1,
            "expected_heading_paths": ["供应商准入与持续评估规范", "2. 标准准入流程", "2.1 客户数据处理的强制条件"],
            "min_chunk_count": 8, "max_chunk_count": 24,
        },
        {
            "document_key": "security-response-pdf", "filename": "security_incident_response.pdf", "file_type": "pdf",
            "expected_parser": "pdf_layout_parser", "expected_page_count": 4, "expected_table_count": 1,
            "expected_element_count": 35, "expected_section_count": 15, "expected_block_count": 35,
            "expected_heading_count": 15, "expected_list_count": 0,
            "expected_block_types": ["heading", "paragraph", "table"],
            "expected_heading_paths": ["2. P1 首要时限", "7. 分级与响应矩阵", "9. 恢复与关闭"],
            "min_chunk_count": 8, "max_chunk_count": 30,
        },
        {
            "document_key": "employee-handbook-docx", "filename": "employee_handbook.docx", "file_type": "docx",
            "expected_parser": "docx_parser", "expected_page_count": 5, "expected_table_count": 2,
            "expected_element_count": 30, "expected_section_count": 15, "expected_block_count": 30,
            "expected_heading_count": 15, "expected_list_count": 6,
            "expected_block_types": ["paragraph", "list", "table"],
            "expected_heading_paths": ["第二章 请假与远程办公", "第三章 信息安全", "第四章 入职、转岗与离职"],
            "min_chunk_count": 8, "max_chunk_count": 24,
        },
        {
            "document_key": "budget-risk-xlsx", "filename": "budget_and_risk_register.xlsx", "file_type": "xlsx",
            "expected_parser": "excel_parser", "expected_sheet_count": 3, "expected_table_count": 4,
            "expected_element_count": 4, "expected_section_count": 3, "expected_block_count": 4,
            "expected_heading_count": 0, "expected_list_count": 0,
            "expected_sheet_names": ["预算总表", "风险清单", "供应商评分"], "expected_block_types": ["table"],
            "expected_heading_paths": ["预算总表", "风险清单", "供应商评分"],
            "min_chunk_count": 4, "max_chunk_count": 20,
        },
    ]
    for spec in specs:
        source_path = SOURCE_DIR / str(spec["filename"])
        spec["sha256"] = sha256(source_path)
        spec["size_bytes"] = source_path.stat().st_size

    manifest = {
        "dataset_id": "multiformat-enterprise-e2e-v1",
        "version": "1.0.0",
        "organization_slug": "org-demo",
        "knowledge_base_name": "企业制度与运营知识库 E2E",
        "license": "Synthetic test data; repository use permitted",
        "source": "Project-authored deterministic fixtures",
        "documents": specs,
    }
    queries = build_queries()
    parser_expectations = {
        "documents": {
            spec["document_key"]: {
                key: value for key, value in spec.items()
                if key.startswith("expected_") and key not in {"expected_heading_paths", "expected_block_types"}
            } | {
                "required_block_types": spec["expected_block_types"],
                "required_heading_paths": spec["expected_heading_paths"],
            }
            for spec in specs
        }
    }
    chunk_expectations = {
        "defaults": {"target_chunk_size": 850, "max_chunk_size": 1000, "chunk_overlap": 200},
        "global_rules": {
            "must_not_start_with_fragment": True,
            "table_chunks_retain_header": True,
            "heading_prefix_matches_metadata": True,
            "required_metadata_any": ["page_start", "sheet_name", "heading_path"],
        },
        "documents": {
            spec["document_key"]: {
                "min_chunk_count": spec["min_chunk_count"],
                "max_chunk_count": spec["max_chunk_count"],
                "required_content_terms": {
                    "travel-reimbursement-txt": ["20,000 元", "发票", "行程单", "审批单"],
                    "supplier-admission-md": ["安全问卷", "隐私影响评估", "customer_data"],
                    "security-response-pdf": ["15 分钟", "30 分钟", "P1"],
                    "employee-handbook-docx": ["远程办公", "客户数据", "账号停用"],
                    "budget-risk-xlsx": ["AI 知识库重构", "135000", "R-003"],
                }[spec["document_key"]],
            }
            for spec in specs
        },
    }
    retrieval_expectations = {
        "query_count": len(queries),
        "category_counts": {
            category: sum(1 for item in queries if item["category"] == category)
            for category in sorted({str(item["category"]) for item in queries})
        },
        "thresholds": {"recall_at_5": 0.8, "mrr": 0.6, "citation_accuracy": 0.8, "no_answer_false_answer_count": 0},
        "answerable_query_ids": [item["query_id"] for item in queries if item["expected_document_keys"]],
        "no_answer_query_ids": [item["query_id"] for item in queries if item["category"] == "no_answer"],
        "permission_query_ids": [item["query_id"] for item in queries if item["category"] == "permission"],
    }
    write_json(FIXTURE_ROOT / "manifest.json", manifest)
    write_json(FIXTURE_ROOT / "queries.json", queries)
    write_json(EXPECTED_DIR / "parser_expectations.json", parser_expectations)
    write_json(EXPECTED_DIR / "chunk_expectations.json", chunk_expectations)
    write_json(EXPECTED_DIR / "retrieval_expectations.json", retrieval_expectations)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--contracts-only", action="store_true")
    args = parser.parse_args()
    SOURCE_DIR.mkdir(parents=True, exist_ok=True)
    EXPECTED_DIR.mkdir(parents=True, exist_ok=True)
    if not args.contracts_only:
        write_text_sources()
        create_pdf()
        create_docx()
    write_contract_files()
    print(f"generated multiformat fixtures in {FIXTURE_ROOT}")


if __name__ == "__main__":
    main()
