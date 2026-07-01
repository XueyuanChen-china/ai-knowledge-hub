from pathlib import Path
import sys


BACKEND_DIR = Path(__file__).resolve().parents[1]
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))


FIXTURE_ROOT = BACKEND_DIR / "tests" / "fixtures" / "splitter_regression"
BINARY_ROOT = FIXTURE_ROOT / "binary_samples"


def generate_excel_sample() -> None:
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Orders"
    sheet.append(["order_id", "customer", "amount", "status"])
    sheet.append([1001, "Alice", 95, "paid"])
    sheet.append([1002, "Bob", 120, "draft"])
    sheet.append([1003, "Carol", 88, "paid"])
    sheet.append([None, None, None, None])

    summary = workbook.create_sheet("Summary")
    summary.append(["metric", "value"])
    summary.append(["total_orders", 3])
    summary.append(["active_status", "paid"])

    workbook.save(BINARY_ROOT / "workbook_orders.xlsx")


def generate_docx_sample() -> None:
    from docx import Document

    document = Document()
    document.add_heading("员工手册", level=1)
    document.add_paragraph("第一章说明系统的录入和审核规则。")
    document.add_heading("提交流程", level=2)
    document.add_paragraph("录入完成后，需要先提交给审核人。")

    list_style = None
    for style_name in ("List Bullet", "列表项目符号", "List Paragraph"):
        try:
            list_style = document.styles[style_name]
            break
        except KeyError:
            continue

    if list_style is not None:
        document.add_paragraph("准备材料", style=list_style)
        document.add_paragraph("提交审批", style=list_style)
    else:
        document.add_paragraph("准备材料")
        document.add_paragraph("提交审批")

    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "说明"
    table.cell(1, 0).text = "status"
    table.cell(1, 1).text = "draft"
    table.cell(2, 0).text = "owner"
    table.cell(2, 1).text = "employee"

    document.save(BINARY_ROOT / "word_handbook.docx")


def generate_pdf_sample() -> None:
    from reportlab.pdfgen import canvas

    pdf_path = BINARY_ROOT / "pdf_policy.pdf"
    pdf_canvas = canvas.Canvas(str(pdf_path), pagesize=(612, 792))

    def draw_header_footer(page_no: int) -> None:
        pdf_canvas.setFont("Helvetica", 10)
        pdf_canvas.drawString(72, 770, "Knowledge Policy")
        pdf_canvas.drawString(72, 24, f"Page {page_no}")

    draw_header_footer(1)
    pdf_canvas.setFont("Helvetica-Bold", 16)
    pdf_canvas.drawString(72, 720, "1. Review Rules")
    pdf_canvas.setFont("Helvetica", 12)
    pdf_canvas.drawString(72, 690, "Reviewers must check title, body, tags, and status.")
    pdf_canvas.drawString(72, 670, "Only active records can enter retrieval.")

    table_x = 72
    table_y_top = 620
    col_widths = [150, 170]
    row_height = 24
    rows = [
        ["field", "value"],
        ["status", "active"],
        ["source_type", "document"],
    ]
    total_width = sum(col_widths)
    total_height = row_height * len(rows)

    for row_index in range(len(rows) + 1):
        y = table_y_top - row_index * row_height
        pdf_canvas.line(table_x, y, table_x + total_width, y)

    current_x = table_x
    pdf_canvas.line(current_x, table_y_top, current_x, table_y_top - total_height)
    for width in col_widths:
        current_x += width
        pdf_canvas.line(current_x, table_y_top, current_x, table_y_top - total_height)

    pdf_canvas.setFont("Helvetica", 11)
    for row_index, row in enumerate(rows):
        text_y = table_y_top - row_height * (row_index + 0.7)
        cell_x = table_x + 8
        for cell_index, cell_text in enumerate(row):
            pdf_canvas.drawString(cell_x, text_y, cell_text)
            cell_x += col_widths[cell_index]

    pdf_canvas.showPage()
    draw_header_footer(2)
    pdf_canvas.setFont("Helvetica-Bold", 16)
    pdf_canvas.drawString(72, 720, "2. Upload Flow")
    pdf_canvas.setFont("Helvetica", 12)
    pdf_canvas.drawString(72, 690, "TXT, Markdown, Word, PDF, and Excel can be parsed.")
    pdf_canvas.drawString(72, 670, "Metadata keeps source and page traceability.")
    pdf_canvas.save()


def main() -> None:
    BINARY_ROOT.mkdir(parents=True, exist_ok=True)
    generate_excel_sample()
    generate_docx_sample()
    generate_pdf_sample()
    print(f"generated binary fixtures in {BINARY_ROOT}")


if __name__ == "__main__":
    main()
