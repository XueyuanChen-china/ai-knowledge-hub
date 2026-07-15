from pathlib import Path
from textwrap import dedent

from docx import Document
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


OUTPUT_DIR = Path(__file__).resolve().parents[1] / "data" / "sample_index_files"


TXT_CONTENT = dedent(
    """
    企业差旅与费用报销操作指引
    版本：2026.07
    发布部门：财务共享中心

    一、适用范围
    本指引适用于总部职能部门、区域销售团队、项目交付团队和外派支持人员。员工因业务出差产生的交通、住宿、餐饮和客户接待费用，均需要按照本指引完成申请、报销与归档。实习生和供应商驻场人员如有特殊安排，以合同和项目补充协议为准。

    二、出差前审批要求
    员工出差前需要至少提前两个工作日在 OA 系统中提交出差申请，写明出差城市、起止日期、业务目标、同行人员和预计预算。部门负责人审批通过后，预算金额会同步到财务台账。若预计单次差旅费用超过 8000 元，还需要区域负责人进行二级审批。

    2.1 预算类特殊要求
    单次出差涉及客户活动、展会布展或现场设备运输时，应在申请单中补充预算拆分说明，明确交通、住宿、物料、接待和应急费用口径。若预计发生外币支付，还需要在备注中写明结算币种、汇率参考口径和付款主体。

    2.2 审批链补充说明
    若出差同行人员来自多个部门，默认由发起人所在部门负责人先审，再由预算归属部门负责人确认。涉及销售折扣、样机借用或渠道返利沟通的差旅，还应同步给对应业务运营负责人留痕。

    三、交通费用标准
    高铁优先选择二等座，单程超过 5 小时且到达时间晚于 22:00 时，可申请一等座。国内航班仅限购买经济舱，国际航班按照岗位等级匹配舱位。网约车仅支持普通快车，特殊情况如深夜抵达、携带大件设备或到达地无公共交通，需要在报销备注中说明原因。

    四、住宿费用标准
    一线城市单晚住宿标准为 550 元，重点客户拜访或展会期间可上浮至 700 元；新一线及省会城市单晚标准为 420 元；其他城市单晚标准为 320 元。若酒店价格超出标准，需要在申请阶段说明原因，并由部门负责人批准后方可报销。

    五、票据与附件要求
    报销时必须上传发票扫描件、行程单、酒店水单和付款截图。若为电子发票，需要确保票面信息清晰，且发票抬头为公司全称。若因特殊情况无法获取原始票据，应补充《票据缺失说明》，由直属上级签字确认后提交财务复核。

    六、常见退回原因
    1. 同一行程包含多个城市时，未按城市拆分填写，导致预算和实际行程无法对应。
    2. 客户接待费用缺少参会人员名单和接待目的说明。
    3. 自驾出行未提供审批记录、公里数说明和油费发票。
    4. 报销单提交后被退回，但员工未在三个工作日内补齐材料重新提交。

    七、流程总结
    整体流程可以概括为：出差申请 -> 审批通过 -> 完成行程 -> 收集票据 -> 提交报销 -> 财务复核 -> 打款归档。若问题集中在票据不完整、预算超标或审批链缺失，财务会先退回单据，再要求补充说明。

    附录A：材料核对清单
    申请阶段：出差目的、预算说明、同行人、审批链。
    报销阶段：发票、行程单、酒店水单、付款截图、差旅总结。
    复核阶段：预算占用记录、例外审批说明、附件是否齐全。
    """
).strip()


MD_CONTENT = dedent(
    """
    # 项目交付知识库维护规范

    > 适用对象：实施、售后、技术支持、项目经理。

    ## 1. 文档目标
    项目交付知识库用于沉淀实施过程中的通用经验、环境配置说明、故障排查方法和上线回滚策略。它既服务新成员快速上手，也服务跨项目复用，避免每次问题都从零排查。

    ## 2. 录入原则
    - 只记录可复用的信息，不记录一次性的闲聊结论。
    - 标题需要明确场景，例如“支付网关切换失败排查”比“问题记录”更好。
    - 正文要尽量包含前置条件、触发症状、排查步骤、最终处理方法和后续预防建议。

    ## 3. 推荐结构
    每篇知识条目推荐使用以下结构：

    ### 3.1 背景
    说明问题发生在什么环境、哪个系统、什么版本，以及问题影响范围。

    ### 3.2 症状
    列出报错表现，例如页面白屏、接口 500、任务队列阻塞、数据库连接数异常升高等。若有日志关键字，也应该直接写进正文，方便后续检索。

    ### 3.3 排查步骤
    1. 先确认配置中心中的开关状态。
    2. 再检查最近一次发布版本和变更清单。
    3. 如果问题涉及数据库，再对比连接池、慢 SQL 和锁等待情况。
    4. 如果问题涉及外部依赖，再检查超时、重试和熔断日志。

    ### 3.4 结论
    给出最终原因和解决方案。例如：因为灰度环境的回调地址未更新，导致支付渠道验签失败；修复方式是更新配置并重新下发证书。

    ### 3.5 示例代码
    如果问题涉及脚本或配置，建议保留最小可复现片段：

    ```bash
    curl -X POST http://127.0.0.1:8000/documents/12/index
    ```

    ```json
    {
      "knowledge_base_id": 7,
      "status": "draft",
      "source_type": "document"
    }
    ```

    ## 4. 审核规则
    知识条目初始状态建议为 `draft`。项目经理或技术负责人确认内容准确后，再切换为 `active`。过时的条目不要直接删除，而是转成 `disabled`，保留历史追溯能力。

    ### 4.1 审核人关注点
    - 内容是否有明确适用范围。
    - 是否写清楚前置条件和失败边界。
    - 是否存在只能在单一项目复用的临时结论。

    ## 5. 示例表格
    | 场景 | 必填字段 | 说明 |
    | --- | --- | --- |
    | 故障排查 | 背景、症状、步骤、结论 | 便于快速复盘 |
    | 上线说明 | 版本、影响范围、回滚方案 | 便于交接 |
    | 接口对接 | 请求参数、鉴权方式、示例响应 | 便于开发联调 |

    ## 6. 额外建议
    如果一篇内容很长，建议在开头补一个摘要；如果同一主题经常被搜索，可以增加别名、关键词和常见问法。例如“支付签约失败”也可以补充“绑卡失败”“扣款协议异常”等近义表达。

    ## 7. FAQ
    ### 7.1 一个问题能拆成多篇吗
    可以。若同一现象背后存在多种根因，建议拆成多篇条目，并在每篇开头写清区分条件，这比把多种路径硬塞进同一个长文档更利于检索和审核。

    ### 7.2 多久回收一次失效知识
    建议按月巡检 active 条目，重点关注最近三个月无人访问、版本明显过旧或依赖已下线系统的内容。
    """
).strip()


PDF_TITLE = "采购与供应商准入管理制度"

PDF_PAGE_ONE_LEFT = [
    (
        "制度目的与适用范围",
        "本制度用于规范采购申请、供应商评估、合同签署、到货验收与付款流程，适用于行政采购、IT 设备采购、营销物料采购和外部服务采购。制度强调在业务效率、成本控制和合规审查之间取得平衡，避免出现先采购后补流程、重复采购或信息不对称的问题。金额较大或涉及核心数据处理能力的供应商，需要增加法务和信息安全联合评审。"
    ),
    (
        "供应商准入流程",
        "业务部门先提出采购需求，采购专员收集至少三家候选供应商信息，并形成比价表。若采购内容涉及系统开发、托管服务或客户数据处理，必须同步完成安全问卷和隐私影响评估。评估通过后，采购经理发起准入审批，法务审核合同条款，财务确认付款条件，再进入下单与履约阶段。"
    ),
]

PDF_PAGE_ONE_RIGHT = [
    (
        "评估维度",
        "评估维度包括四类：一是交付能力，包括实施经验、项目经理稳定性和交付周期；二是成本与付款条款，包括是否支持分阶段付款、是否有预付款要求；三是安全与合规，包括日志留存、数据隔离、人员权限和漏洞修复时效；四是服务保障，包括响应 SLA、升级机制和驻场支持能力。"
    ),
    (
        "检查项摘要",
        "1. 营业执照、行业资质、近三年案例是否齐全。 2. 权限模型、日志留存、漏洞修复 SLA 是否通过安全问卷复核。 3. 数据处理条款、保密约定、分包限制是否在合同中明确。 4. 项目经理安排、升级机制、值守支持是否覆盖关键场景。"
    ),
]

PDF_PAGE_TWO_LEFT = [
    (
        "高风险信号",
        "若供应商无法提供历史客户案例、无法说明权限管理方案，或合同中拒绝承担数据泄露责任，应视为高风险。若单次采购金额超过二十万元，需要由采购委员会复核，并在立项会上同步说明预算依据。若存在跨境数据处理、重要系统接入或长期驻场开发，还应补充更细的专项审查。"
    ),
    (
        "审批记录要求",
        "每次准入审批都应在系统中保留申请人、审批链、比价附件、风险意见和最终决策。若审批过程中有例外放行，必须明确例外原因、补救动作和跟踪人，避免后续复盘时缺少依据。"
    ),
]

PDF_PAGE_TWO_RIGHT = [
    (
        "常见复盘问题",
        "采购复盘中常见问题包括：供应商名单收集不完整、比价说明缺少口径、合同版本不统一、交付节点未写入验收条件，以及安全整改项关闭状态没有回写到采购台账。以上问题会直接影响后续验收、付款和审计追溯。"
    ),
    (
        "执行建议",
        "建议把供应商分类、数据访问等级、合同模板版本、付款节点和整改承诺统一写入知识库，并与采购表单联动。这样后续做语义搜索时，不仅能搜到制度原文，也能搜到对应的执行口径和风险提示。"
    ),
]

PDF_SUMMARY_ROWS = [
    ["检查项", "要求"],
    ["安全问卷", "涉及系统开发或数据处理时必须完成"],
    ["法务审核", "合同条款、保密义务、分包限制必须审阅"],
    ["采购复核", "单次采购金额超过二十万元需委员会复核"],
]


DOCX_TABLE_ROWS = [
    ("阶段", "负责人", "关键输出", "时限"),
    ("需求提交", "业务部门", "采购申请单、预算说明", "T+0"),
    ("供应商初筛", "采购专员", "候选名单、基础比价", "T+2"),
    ("风险评估", "信息安全/法务", "安全问卷、合同审阅意见", "T+5"),
    ("审批与下单", "采购经理", "审批记录、订单信息", "T+7"),
]

DOCX_RISK_ROWS = [
    ("问题类型", "触发信号", "建议动作"),
    ("交接缺失", "客户目标、关键联系人未补齐", "退回交接单并补充模板字段"),
    ("知识分散", "同一主题出现多个命名不一致条目", "统一标题规范并合并别名"),
    ("升级缓慢", "跨部门问题 30 分钟内无人认领", "建立升级群和责任人轮值"),
]


XLSX_ROWS = {
    "预算总表": [
        ["部门", "项目", "季度", "预算金额", "已使用金额", "负责人", "备注"],
        ["销售一部", "华东渠道拓展", "2026Q3", "180000", "62500", "李晨", "含展会与客户活动"],
        ["交付中心", "重点项目驻场支持", "2026Q3", "96000", "24500", "周楠", "含差旅与加班餐补"],
        ["产品部", "AI 知识库重构", "2026Q3", "135000", "48200", "王璐", "含模型调用与外包测试"],
        ["行政部", "办公室设备更新", "2026Q3", "68000", "12000", "陈希", "优先更换会议室设备"],
        ["财务部", "报销流程电子化", "2026Q3", "45000", "8000", "宋妍", "预计 8 月启动"],
        [],
        ["审批统计", "数量"],
        ["待审批申请", "12"],
        ["已退回申请", "4"],
    ],
    "风险清单": [
        ["风险编号", "风险描述", "影响等级", "触发条件", "应对措施", "责任人"],
        ["R-001", "预算超支导致项目延期", "高", "连续两月消耗超预算 20%", "冻结非核心支出并重排里程碑", "李晨"],
        ["R-002", "供应商交付延迟", "中", "关键里程碑延期超过 5 天", "启动备选供应商并拆分交付范围", "周楠"],
        ["R-003", "知识库数据权限配置错误", "高", "非授权成员可访问受限条目", "按权限组复核并补审计日志", "王璐"],
        ["R-004", "上线窗口与客户变更冲突", "中", "客户冻结期提前开始", "提前准备回滚方案并调整发布节奏", "陈希"],
    ],
    "供应商评分": [
        ["供应商", "交付能力", "成本", "安全合规", "服务保障", "结论"],
        ["云启科技", "4.5", "4.0", "4.8", "4.2", "建议准入"],
        ["星河系统", "4.2", "4.6", "3.9", "4.0", "需补安全材料"],
        ["同舟数据", "3.8", "4.3", "4.7", "4.5", "可进入复审"],
    ],
}


def write_text_file(path: Path, content: str) -> None:
    path.write_text(content + "\n", encoding="utf-8")


def pick_docx_style(document: Document, candidates: list[str]):
    for style_name in candidates:
        try:
            return document.styles[style_name]
        except KeyError:
            continue
    return None


def create_docx_file(path: Path) -> None:
    document = Document()
    document.add_heading("客户成功团队季度复盘与改进计划", level=1)

    sections = [
        (
            "一、整体情况",
            "本季度客户成功团队重点支持了金融、零售和制造三个行业的 12 个重点客户。整体续约率保持在 91%，但高价值客户的功能培训参与度下降，导致部分客户在高级能力的使用深度上没有达到预期。",
        ),
        (
            "二、主要问题",
            "问题主要集中在三个方面：第一，实施交接资料不够完整，客户成功经理接手后需要重新确认背景；第二，知识库中关于权限配置和报表异常的内容较分散，检索效率偏低；第三，跨部门升级问题的响应链条较长，影响客户满意度。",
        ),
        (
            "三、改进动作",
            "计划从交接模板、知识沉淀、升级机制三条线并行推进。交接模板增加客户目标、关键联系人、已交付范围和风险项；知识沉淀统一标题规范与标签标准；升级机制要求在 30 分钟内完成责任人确认。",
        ),
        (
            "四、补充说明",
            "对于处于续约前 60 天的客户，需要单独建立风险跟踪列表，并在周会上同步使用情况、待解决问题和下一步动作。若客户提出新的安全或合规要求，应在一个工作日内同步给产品与法务团队。",
        ),
    ]

    for heading, paragraph in sections:
        document.add_heading(heading, level=2)
        document.add_paragraph(paragraph)

    document.add_heading("三.1 重点动作拆解", level=3)
    document.add_paragraph("为了避免复盘结论停留在口头层面，本季度要求每个改进动作都映射到负责人、截止时间和复核方式。")

    bullet_style = pick_docx_style(document, ["List Bullet", "列表项目符号", "List Paragraph"])
    for item in [
        "统一交接模板字段，新增客户目标、已交付范围和风险项。",
        "补齐知识库标签规范，明确系统名、问题类型和影响范围。",
        "建立升级响应值班表，确保高优先级问题 30 分钟内有人认领。",
    ]:
        if bullet_style is not None:
            document.add_paragraph(item, style=bullet_style)
        else:
            document.add_paragraph(f"- {item}")

    number_style = pick_docx_style(document, ["List Number", "列表编号", "List Paragraph"])
    for step in [
        "每周盘点新增问题单，筛选是否需要沉淀为知识条目。",
        "每月抽查 active 条目，确认是否仍适配当前系统版本。",
        "季度复盘时统一回收失效条目，并输出替代知识链接。",
    ]:
        if number_style is not None:
            document.add_paragraph(step, style=number_style)
        else:
            document.add_paragraph(f"1. {step}")

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, value in enumerate(DOCX_TABLE_ROWS[0]):
        header_cells[index].text = value

    for row in DOCX_TABLE_ROWS[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)

    document.add_paragraph("以下风险对照表用于测试 Word 表格和标题间混排场景。")

    risk_table = document.add_table(rows=1, cols=3)
    risk_table.style = "Table Grid"
    risk_header_cells = risk_table.rows[0].cells
    for index, value in enumerate(DOCX_RISK_ROWS[0]):
        risk_header_cells[index].text = value

    for row in DOCX_RISK_ROWS[1:]:
        cells = risk_table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)

    document.add_paragraph(
        "结论：后续知识库建设应重点覆盖权限配置、客户升级流程、报表异常排查和续约风险判断这四类高频主题。"
    )
    document.save(path)


def create_pdf_file(path: Path) -> None:
    font_path = find_chinese_font_path()
    font_name = "SampleChineseFont"
    pdfmetrics.registerFont(TTFont(font_name, str(font_path)))

    pdf = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    left_margin = 50
    right_margin = 50
    top = height - 58
    bottom = 56
    gutter = 40
    column_width = (width - left_margin - right_margin - gutter) / 2
    line_height = 15
    paragraph_gap = 10

    pdf.setTitle(PDF_TITLE)

    def draw_page_header_footer(page_number: int) -> None:
        pdf.setFont(font_name, 9)
        pdf.drawString(left_margin, height - 28, "采购制度样例文档")
        pdf.drawRightString(width - right_margin, 28, f"Page {page_number}")

    def wrap_text(text: str, *, max_width: float, font_size: int) -> list[str]:
        normalized = " ".join(text.strip().split())
        if not normalized:
            return []

        chunks: list[str] = []
        current = ""
        for char in normalized:
            tentative = current + char
            tentative_width = pdf.stringWidth(tentative, font_name, font_size)
            if tentative_width > max_width and current:
                chunks.append(current)
                current = char
            else:
                current = tentative
        if current:
            chunks.append(current)
        return chunks

    def draw_column_section(
        x: float,
        y: float,
        heading: str,
        body: str,
    ) -> float:
        pdf.setFont(font_name, 13)
        pdf.drawString(x, y, heading)
        y -= line_height

        pdf.setFont(font_name, 10)
        for line in wrap_text(body, max_width=column_width - 6, font_size=10):
            pdf.drawString(x, y, line)
            y -= line_height

        return y - paragraph_gap

    def draw_summary_table(x: float, y_top: float) -> float:
        col_widths = [72, column_width - 72]
        total_width = sum(col_widths)
        wrapped_rows: list[list[list[str]]] = []
        row_heights: list[float] = []

        for row in PDF_SUMMARY_ROWS:
            wrapped_cells: list[list[str]] = []
            max_line_count = 1
            for cell_index, cell_text in enumerate(row):
                cell_lines = wrap_text(
                    cell_text,
                    max_width=col_widths[cell_index] - 8,
                    font_size=9,
                ) or [""]
                wrapped_cells.append(cell_lines)
                max_line_count = max(max_line_count, len(cell_lines))
            wrapped_rows.append(wrapped_cells)
            row_heights.append(max(20, max_line_count * 11 + 8))

        total_height = sum(row_heights)
        bottom_y = y_top - total_height

        pdf.setLineWidth(0.8)
        current_y = y_top
        pdf.line(x, current_y, x + total_width, current_y)
        for row_height in row_heights:
            current_y -= row_height
            pdf.line(x, current_y, x + total_width, current_y)

        current_x = x
        pdf.line(current_x, y_top, current_x, bottom_y)
        for width_part in col_widths:
            current_x += width_part
            pdf.line(current_x, y_top, current_x, bottom_y)

        pdf.setFont(font_name, 9)
        current_y = y_top
        for row_index, wrapped_row in enumerate(wrapped_rows):
            row_height = row_heights[row_index]
            text_top_y = current_y - 12
            current_x = x
            for cell_index, cell_lines in enumerate(wrapped_row):
                line_y = text_top_y
                for cell_line in cell_lines:
                    pdf.drawString(current_x + 4, line_y, cell_line)
                    line_y -= 10
                current_x += col_widths[cell_index]
            current_y -= row_height

        return bottom_y - paragraph_gap

    page_number = 1
    draw_page_header_footer(page_number)

    pdf.setFont(font_name, 16)
    pdf.drawString(left_margin, top, PDF_TITLE)

    left_x = left_margin
    right_x = left_margin + column_width + gutter
    column_start_y = top - 32

    left_y = column_start_y
    for heading, body in PDF_PAGE_ONE_LEFT:
        left_y = draw_column_section(left_x, left_y, heading, body)

    right_y = column_start_y - 6
    for heading, body in PDF_PAGE_ONE_RIGHT:
        right_y = draw_column_section(right_x, right_y, heading, body)

    pdf.showPage()
    page_number += 1
    draw_page_header_footer(page_number)

    pdf.setFont(font_name, 14)
    pdf.drawString(left_margin, top, "准入复核与执行建议")

    column_start_y = top - 28
    left_y = column_start_y
    for heading, body in PDF_PAGE_TWO_LEFT:
        left_y = draw_column_section(left_x, left_y, heading, body)

    right_y = column_start_y - 6
    for heading, body in PDF_PAGE_TWO_RIGHT:
        right_y = draw_column_section(right_x, right_y, heading, body)
    right_y = draw_summary_table(right_x, right_y)

    pdf.save()


def find_chinese_font_path() -> Path:
    candidates = [
        Path("/Library/Fonts/Arial Unicode.ttf"),
        Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
        Path("/System/Library/Fonts/Hiragino Sans GB.ttc"),
        Path("/System/Library/Fonts/STHeiti Medium.ttc"),
        Path("/System/Library/Fonts/Supplemental/Songti.ttc"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    raise RuntimeError("no suitable Chinese font found for sample pdf generation")


def create_xlsx_file(path: Path) -> None:
    workbook = Workbook()
    default_sheet = workbook.active
    workbook.remove(default_sheet)

    for sheet_name, rows in XLSX_ROWS.items():
        sheet = workbook.create_sheet(title=sheet_name)
        for row in rows:
            sheet.append(row)

        for column in sheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                value = "" if cell.value is None else str(cell.value)
                max_length = max(max_length, len(value))
            sheet.column_dimensions[column_letter].width = min(max_length + 4, 28)

    workbook.save(path)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    write_text_file(OUTPUT_DIR / "sample_policy_notice.txt", TXT_CONTENT)
    write_text_file(OUTPUT_DIR / "sample_project_knowledge.md", MD_CONTENT)
    create_docx_file(OUTPUT_DIR / "sample_customer_success_review.docx")
    create_pdf_file(OUTPUT_DIR / "sample_supplier_management_policy.pdf")
    create_xlsx_file(OUTPUT_DIR / "sample_budget_and_risk_register.xlsx")

    print(f"Generated sample files in: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
